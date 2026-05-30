# -*- coding: utf-8 -*-
"""
Xuat bang tong quan testcase cua he thong quan ly thuc tap ra file Word.
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Màu sắc ──────────────────────────────────────────────────────────────────
C_HEADER    = RGBColor(0x1F, 0x49, 0x7D)   # xanh đậm
C_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
C_MODULE    = RGBColor(0x2E, 0x74, 0xB5)   # xanh section
C_MODULE_BG = RGBColor(0xD6, 0xE4, 0xF7)   # nền section
C_ODD       = RGBColor(0xED, 0xF3, 0xFB)   # hàng lẻ
C_PASS      = RGBColor(0x00, 0x70, 0x00)
C_CAPTION   = RGBColor(0x1F, 0x49, 0x7D)

# ── Dữ liệu testcase ─────────────────────────────────────────────────────────
# (stt, ma_tc, chuc_nang, mo_ta, du_lieu_vao, ket_qua_mong_doi, trang_thai)
TESTCASES = [
    # ─── MODULE 1: XÁC THỰC (Authentication) ────────────────────────────────
    ("MODULE", "1", "XÁC THỰC HỆ THỐNG (Authentication)", "", "", "", ""),
    (1,"TC001","Đăng nhập – Sinh viên hợp lệ",
     "Đăng nhập bằng tài khoản sinh viên đúng thông tin",
     "userId: SV001 | password: 123456 | role: sinh-vien",
     "Trả về token JWT, thông tin sinh viên, HTTP 200","Pass"),
    (2,"TC002","Đăng nhập – Sai mật khẩu",
     "Đăng nhập với mật khẩu không chính xác",
     "userId: SV001 | password: sai | role: sinh-vien",
     "HTTP 401, message: 'Mật khẩu không chính xác'","Pass"),
    (3,"TC003","Đăng nhập – Sai vai trò",
     "Đăng nhập đúng userId nhưng sai role",
     "userId: GV001 | password: 123456 | role: sinh-vien",
     "HTTP 401, thông báo lỗi vai trò không khớp","Pass"),
    (4,"TC004","Đăng nhập – Tài khoản bị khóa",
     "Đăng nhập bằng tài khoản is_active = 0",
     "userId: SV_LOCKED | password: 123456",
     "HTTP 403, message: 'Tài khoản đã bị vô hiệu hóa'","Pass"),
    (5,"TC005","Đăng nhập – Giảng viên hợp lệ",
     "Đăng nhập bằng tài khoản giảng viên",
     "userId: GV001 | password: 123456 | role: giang-vien",
     "HTTP 200, token JWT, thông tin giảng viên","Pass"),
    (6,"TC006","Đăng nhập – Admin hợp lệ",
     "Đăng nhập bằng tài khoản quản trị viên",
     "userId: admin001 | password: 123456 | role: admin",
     "HTTP 200, token JWT, thông tin admin","Pass"),
    (7,"TC007","Đăng xuất hệ thống",
     "Gọi API logout với token hợp lệ",
     "Bearer token hợp lệ",
     "HTTP 200, message: 'Đăng xuất thành công'","Pass"),
    (8,"TC008","Lấy thông tin người dùng hiện tại",
     "Gọi /auth/me với token hợp lệ",
     "Bearer token hợp lệ",
     "HTTP 200, trả về id, userId, role, name, email","Pass"),
    (9,"TC009","Gọi API không có token",
     "Truy cập API cần xác thực mà không gửi token",
     "Không có Authorization header",
     "HTTP 401, message: 'Unauthorized'","Pass"),

    # ─── MODULE 2: QUẢN LÝ TÀI KHOẢN ────────────────────────────────────────
    ("MODULE", "2", "QUẢN LÝ TÀI KHOẢN (Accounts)", "", "", "", ""),
    (10,"TC010","Tạo tài khoản mới",
     "Admin tạo tài khoản cho sinh viên mới",
     "userId: SV100, email: sv100@test.com, password: 123456, role: sinh-vien",
     "HTTP 201, tài khoản được tạo thành công","Pass"),
    (11,"TC011","Tạo tài khoản – userId đã tồn tại",
     "Tạo tài khoản với userId trùng lặp",
     "userId: SV001 (đã tồn tại)",
     "HTTP 400, message: 'userId đã tồn tại'","Pass"),
    (12,"TC012","Lấy danh sách tài khoản",
     "Admin lấy toàn bộ danh sách tài khoản",
     "GET /accounts với token admin",
     "HTTP 200, danh sách tài khoản có phân trang","Pass"),
    (13,"TC013","Kích hoạt tài khoản",
     "Admin kích hoạt lại tài khoản đã bị khóa",
     "POST /accounts/:userId/activate",
     "HTTP 200, is_active = 1","Pass"),
    (14,"TC014","Vô hiệu hóa tài khoản",
     "Admin vô hiệu hóa tài khoản sinh viên",
     "POST /accounts/:userId/deactivate",
     "HTTP 200, is_active = 0","Pass"),
    (15,"TC015","Đổi mật khẩu",
     "Người dùng đổi mật khẩu cũ sang mới",
     "oldPassword: 123456, newPassword: newPass@123",
     "HTTP 200, mật khẩu được cập nhật","Pass"),
    (16,"TC016","Quên mật khẩu – gửi email",
     "Người dùng yêu cầu reset mật khẩu qua email",
     "email: sv001@dainam.edu.vn",
     "HTTP 200, email chứa mã xác nhận được gửi đi","Pass"),
    (17,"TC017","Xác minh mã reset mật khẩu",
     "Người dùng nhập mã xác nhận hợp lệ",
     "email: sv001@..., code: 123456",
     "HTTP 200, cho phép đặt mật khẩu mới","Pass"),
    (18,"TC018","Reset mật khẩu mới",
     "Người dùng đặt mật khẩu mới sau xác minh",
     "email, code hợp lệ, newPassword: newPass@123",
     "HTTP 200, mật khẩu được cập nhật thành công","Pass"),

    # ─── MODULE 3: QUẢN LÝ SINH VIÊN ────────────────────────────────────────
    ("MODULE", "3", "QUẢN LÝ SINH VIÊN (SinhVien)", "", "", "", ""),
    (19,"TC019","Lấy danh sách sinh viên",
     "Admin/Giảng viên lấy danh sách tất cả sinh viên",
     "GET /sinh-vien | token hợp lệ",
     "HTTP 200, danh sách sinh viên kèm thông tin","Pass"),
    (20,"TC020","Lấy thông tin 1 sinh viên",
     "Lấy chi tiết sinh viên theo id",
     "GET /sinh-vien/:id",
     "HTTP 200, thông tin đầy đủ sinh viên","Pass"),
    (21,"TC021","Thêm sinh viên mới",
     "Admin thêm sinh viên vào hệ thống",
     "ma_sinh_vien, ho_ten, lop, khoa, account_id",
     "HTTP 201, sinh viên được tạo","Pass"),
    (22,"TC022","Cập nhật thông tin sinh viên",
     "Sinh viên/Admin cập nhật hồ sơ",
     "PUT /sinh-vien/:id | các trường cần cập nhật",
     "HTTP 200, thông tin được cập nhật","Pass"),
    (23,"TC023","Xóa sinh viên",
     "Admin xóa sinh viên khỏi hệ thống",
     "DELETE /sinh-vien/:id",
     "HTTP 200, sinh viên bị xóa","Pass"),
    (24,"TC024","Tải lên CV của sinh viên",
     "Sinh viên upload file CV (PDF/DOCX)",
     "Multipart form-data, file CV ≤ 5MB",
     "HTTP 200, cv_path được lưu vào CSDL","Pass"),
    (25,"TC025","Tải lên CV – sai định dạng",
     "Sinh viên upload file không phải PDF/DOCX",
     "File .exe hoặc .jpg",
     "HTTP 400, message: 'Định dạng file không hợp lệ'","Pass"),
    (26,"TC026","Phân tích CV bằng AI",
     "Sinh viên phân tích CV tự động bằng AI",
     "POST /sinh-vien/analyze-cv | cv_path hợp lệ",
     "HTTP 200, kết quả phân tích kỹ năng từ CV","Pass"),
    (27,"TC027","Thống kê sinh viên",
     "Lấy thống kê tổng quan sinh viên trong hệ thống",
     "GET /sinh-vien/stats",
     "HTTP 200, tổng số SV, SV đã phân công, chưa phân công","Pass"),

    # ─── MODULE 4: QUẢN LÝ GIẢNG VIÊN ──────────────────────────────────────
    ("MODULE", "4", "QUẢN LÝ GIẢNG VIÊN (GiangVien)", "", "", "", ""),
    (28,"TC028","Lấy danh sách giảng viên",
     "Admin lấy danh sách tất cả giảng viên",
     "GET /giang-vien | token admin",
     "HTTP 200, danh sách giảng viên","Pass"),
    (29,"TC029","Thêm giảng viên mới",
     "Admin thêm giảng viên vào hệ thống",
     "ma_giang_vien, ho_ten, khoa, bo_mon, account_id",
     "HTTP 201, giảng viên được tạo","Pass"),
    (30,"TC030","Cập nhật thông tin giảng viên",
     "Admin/GV cập nhật hồ sơ giảng viên",
     "PUT /giang-vien/:id | trường cần cập nhật",
     "HTTP 200, thông tin được lưu","Pass"),
    (31,"TC031","Tự động phân công giảng viên",
     "Admin kích hoạt tự động phân công GV cho SV",
     "POST /giang-vien/auto-assign",
     "HTTP 200, số lượng SV được phân công GV","Pass"),
    (32,"TC032","Xuất danh sách giảng viên Excel",
     "Admin xuất danh sách giảng viên ra file Excel",
     "GET /giang-vien/export",
     "HTTP 200, file Excel được tải về","Pass"),

    # ─── MODULE 5: QUẢN LÝ DOANH NGHIỆP ────────────────────────────────────
    ("MODULE", "5", "QUẢN LÝ DOANH NGHIỆP (DoanhNghiep)", "", "", "", ""),
    (33,"TC033","Lấy danh sách doanh nghiệp",
     "Admin lấy danh sách doanh nghiệp trong hệ thống",
     "GET /doanh-nghiep | token admin",
     "HTTP 200, danh sách doanh nghiệp","Pass"),
    (34,"TC034","Thêm doanh nghiệp mới",
     "Admin thêm doanh nghiệp hợp tác mới",
     "ma_doanh_nghiep, ten_cong_ty, dia_chi, so_dien_thoai",
     "HTTP 201, doanh nghiệp được tạo","Pass"),
    (35,"TC035","Cập nhật thông tin doanh nghiệp",
     "Admin/DN cập nhật thông tin doanh nghiệp",
     "PUT /doanh-nghiep/:id | trường cần cập nhật",
     "HTTP 200, thông tin được cập nhật","Pass"),
    (36,"TC036","Xóa doanh nghiệp",
     "Admin xóa doanh nghiệp khỏi hệ thống",
     "DELETE /doanh-nghiep/:id",
     "HTTP 200, doanh nghiệp bị xóa","Pass"),
    (37,"TC037","Xuất danh sách doanh nghiệp",
     "Admin xuất danh sách doanh nghiệp ra Excel",
     "GET /doanh-nghiep/export",
     "HTTP 200, file Excel được tải về","Pass"),

    # ─── MODULE 6: ĐĂNG KÝ THỰC TẬP ────────────────────────────────────────
    ("MODULE", "6", "ĐĂNG KÝ THỰC TẬP (Registration)", "", "", "", ""),
    (38,"TC038","Sinh viên đăng ký thực tập – Khoa giới thiệu",
     "SV đăng ký thực tập theo hình thức khoa giới thiệu",
     "nguyen_vong: khoa-gioi-thieu, vi_tri_thuc_tap_mong_muon",
     "HTTP 201, đơn đăng ký được tạo, trạng thái: CHO_DUYET","Pass"),
    (39,"TC039","Sinh viên đăng ký thực tập – Tự liên hệ",
     "SV đăng ký với công ty tự tìm được",
     "nguyen_vong: tu-lien-he, ten_cong_ty, dia_chi_cong_ty",
     "HTTP 201, đơn đăng ký được tạo với thông tin công ty","Pass"),
    (40,"TC040","Kiểm tra trạng thái đăng ký",
     "SV kiểm tra trạng thái đơn đăng ký hiện tại",
     "GET /registration/check | token SV",
     "HTTP 200, trạng thái và thông tin đơn đăng ký","Pass"),
    (41,"TC041","Admin duyệt đăng ký thực tập",
     "Admin phê duyệt đơn đăng ký của sinh viên",
     "PUT /admin/student-registrations/:id/approve",
     "HTTP 200, trạng thái chuyển thành DA_DUYET","Pass"),
    (42,"TC042","Admin từ chối đăng ký thực tập",
     "Admin từ chối đơn với lý do cụ thể",
     "PUT /admin/student-registrations/:id/reject | ly_do_tu_choi",
     "HTTP 200, trạng thái: TU_CHOI, lý do được lưu","Pass"),
    (43,"TC043","Duyệt hàng loạt theo nguyện vọng",
     "Admin duyệt hàng loạt SV theo nguyện vọng",
     "POST /admin/students/bulk-approve-by-preference",
     "HTTP 200, số lượng SV được duyệt","Pass"),

    # ─── MODULE 7: QUY TRÌNH PHỎNG VẤN ─────────────────────────────────────
    ("MODULE", "7", "QUY TRÌNH PHỎNG VẤN (Interview Workflow)", "", "", "", ""),
    (44,"TC044","SV nộp đơn ứng tuyển vào DN",
     "Sinh viên ứng tuyển vào vị trí thực tập của DN",
     "POST /interview-workflow/student/applications | tin_tuyen_dung_id, thu_xin_viec",
     "HTTP 201, đơn ứng tuyển được tạo, trạng thái: PENDING","Pass"),
    (45,"TC045","Xem danh sách đơn ứng tuyển – SV",
     "Sinh viên xem lịch sử các đơn ứng tuyển",
     "GET /interview-workflow/student/applications",
     "HTTP 200, danh sách đơn ứng tuyển của SV","Pass"),
    (46,"TC046","Admin xem và duyệt đơn ứng tuyển",
     "Admin xem toàn bộ đơn ứng tuyển và duyệt",
     "PUT /interview-workflow/admin/applications/:id/review | action: APPROVED",
     "HTTP 200, trạng thái: APPROVED","Pass"),
    (47,"TC047","Doanh nghiệp lên lịch phỏng vấn",
     "DN đặt lịch phỏng vấn cho SV đã được duyệt",
     "PUT /interview-workflow/company/applications/:id/interview | date, time, location",
     "HTTP 200, lịch phỏng vấn được lưu, trạng thái: INTERVIEW_SCHEDULED","Pass"),
    (48,"TC048","Doanh nghiệp cập nhật kết quả phỏng vấn",
     "DN nhập kết quả phỏng vấn (đậu/rớt)",
     "PUT /interview-workflow/company/applications/:id/result | result: PASS/FAIL",
     "HTTP 200, kết quả được lưu","Pass"),

    # ─── MODULE 8: PHÂN CÔNG THỰC TẬP ──────────────────────────────────────
    ("MODULE", "8", "PHÂN CÔNG THỰC TẬP (Assignments)", "", "", "", ""),
    (49,"TC049","Tạo phân công thực tập",
     "Admin tạo phân công SV vào DN và GV hướng dẫn",
     "sinh_vien_id, doanh_nghiep_id, giang_vien_id, dot_thuc_tap_id, ngay_bat_dau, ngay_ket_thuc",
     "HTTP 201, phân công được tạo","Pass"),
    (50,"TC050","Lấy danh sách phân công",
     "Admin/GV xem danh sách phân công thực tập",
     "GET /assignments | token hợp lệ",
     "HTTP 200, danh sách phân công","Pass"),
    (51,"TC051","Tự động phân công theo khoa giới thiệu",
     "Admin kích hoạt tự động phân công DN cho SV khoa giới thiệu",
     "POST /admin/students/auto-assign-company-khoa-gioi-thieu",
     "HTTP 200, số lượng SV được phân công DN","Pass"),
    (52,"TC052","Xóa phân công thực tập",
     "Admin hủy phân công thực tập",
     "DELETE /assignments/:id",
     "HTTP 200, phân công bị xóa","Pass"),

    # ─── MODULE 9: ĐỢT THỰC TẬP ─────────────────────────────────────────────
    ("MODULE", "9", "ĐỢT THỰC TẬP (InternshipBatches)", "", "", "", ""),
    (53,"TC053","Tạo đợt thực tập mới",
     "Admin tạo đợt thực tập mới",
     "ten_dot, thoi_gian_bat_dau, thoi_gian_ket_thuc, mo_ta",
     "HTTP 201, đợt thực tập được tạo, trạng thái: sap-mo","Pass"),
    (54,"TC054","Lấy danh sách đợt thực tập",
     "Xem tất cả đợt thực tập trong hệ thống",
     "GET /internship-batches",
     "HTTP 200, danh sách các đợt thực tập","Pass"),
    (55,"TC055","Lấy đợt thực tập đang diễn ra",
     "Lấy thông tin đợt thực tập đang active",
     "GET /internship-batches/active",
     "HTTP 200, thông tin đợt đang dien ra","Pass"),
    (56,"TC056","Import danh sách tham gia từ Excel",
     "Admin import danh sách SV/GV/DN tham gia đợt từ file Excel",
     "POST /internship-batches/:id/import-participants | file Excel",
     "HTTP 200, số lượng bản ghi được import","Pass"),
    (57,"TC057","Xuất danh sách tham gia đợt",
     "Admin xuất danh sách tham gia đợt thực tập ra Excel",
     "GET /internship-batches/:id/export",
     "HTTP 200, file Excel được tải về","Pass"),
    (58,"TC058","Xóa đợt thực tập",
     "Admin xóa đợt thực tập chưa có dữ liệu",
     "DELETE /internship-batches/:id",
     "HTTP 200, đợt bị xóa","Pass"),

    # ─── MODULE 10: BÁO CÁO THỰC TẬP ───────────────────────────────────────
    ("MODULE", "10", "BÁO CÁO THỰC TẬP (Reports)", "", "", "", ""),
    (59,"TC059","Sinh viên nộp báo cáo tuần",
     "SV tạo báo cáo thực tập theo tuần",
     "phan_cong_id, loai_bao_cao: tuan, tieu_de, noi_dung, ngay_nop",
     "HTTP 201, báo cáo được tạo, trạng thái: chua-duyet","Pass"),
    (60,"TC060","Sinh viên nộp báo cáo cuối kỳ",
     "SV nộp báo cáo tổng kết cuối đợt thực tập",
     "phan_cong_id, loai_bao_cao: cuoi-khoa, file đính kèm",
     "HTTP 201, báo cáo cuối kỳ được tạo","Pass"),
    (61,"TC061","Giảng viên chấm điểm báo cáo",
     "GV chấm điểm và nhận xét báo cáo của SV",
     "POST /reports/weekly/:id/grade | diem, nhan_xet_gv",
     "HTTP 200, điểm và nhận xét được lưu, trạng thái: da-duyet","Pass"),
    (62,"TC062","GV yêu cầu SV sửa báo cáo",
     "GV trả lại báo cáo yêu cầu sửa",
     "PUT /reports/weekly/:id | trang_thai: can-sua, nhan_xet",
     "HTTP 200, trạng thái: can-sua, SV nhận thông báo","Pass"),
    (63,"TC063","Xem thống kê báo cáo",
     "Xem thống kê tổng quan các báo cáo trong hệ thống",
     "GET /reports/statistics",
     "HTTP 200, số lượng báo cáo theo trạng thái","Pass"),

    # ─── MODULE 11: ĐỢT NỘP & BÀI NỘP ─────────────────────────────────────
    ("MODULE", "11", "ĐỢT NỘP BÁO CÁO & BÀI NỘP (Submissions)", "", "", "", ""),
    (64,"TC064","Giảng viên tạo đợt nộp báo cáo",
     "GV tạo lịch nộp báo cáo theo tuần cho SV",
     "POST /teacher/submissions/slots | tieu_de, loai_bao_cao, start_at, end_at",
     "HTTP 201, đợt nộp được tạo","Pass"),
    (65,"TC065","SV nộp bài trong đợt nộp",
     "Sinh viên upload file báo cáo trong thời hạn",
     "POST /student/slots/:slotId/upload | file PDF/DOCX",
     "HTTP 201, bài nộp được lưu, trạng thái: da_nop","Pass"),
    (66,"TC066","SV nộp bài – ngoài thời hạn",
     "SV cố nộp bài sau khi đợt nộp đã đóng",
     "POST /student/slots/:slotId/upload | sau end_at",
     "HTTP 400, message: 'Đã hết thời gian nộp bài'","Pass"),
    (67,"TC067","GV nhận xét bài nộp",
     "GV duyệt và nhận xét bài nộp của SV",
     "POST /submissions/:submissionId/review | teacher_comment, trang_thai",
     "HTTP 200, nhận xét được lưu","Pass"),
    (68,"TC068","Xem trạng thái nộp của SV trong slot",
     "GV xem ai đã nộp, chưa nộp trong một đợt",
     "GET /teacher/submissions/slots/:slotId/statuses",
     "HTTP 200, danh sách SV và trạng thái nộp","Pass"),
    (69,"TC069","SV xem bài nộp của mình",
     "SV xem lại lịch sử các bài đã nộp",
     "GET /student/all-my-submissions",
     "HTTP 200, danh sách bài nộp của SV","Pass"),
    (70,"TC070","Xuất bảng điểm đợt nộp",
     "GV xuất bảng điểm của đợt nộp ra Excel",
     "GET /report-batches/:id/export",
     "HTTP 200, file Excel bảng điểm","Pass"),

    # ─── MODULE 12: ĐÁNH GIÁ THỰC TẬP ──────────────────────────────────────
    ("MODULE", "12", "ĐÁNH GIÁ THỰC TẬP (Evaluations)", "", "", "", ""),
    (71,"TC071","Doanh nghiệp đánh giá sinh viên",
     "DN chấm điểm và nhận xét sinh viên thực tập",
     "PUT /company-internships/students/:studentId/evaluation | diem_so, nhan_xet",
     "HTTP 200, đánh giá được lưu","Pass"),
    (72,"TC072","DN nộp tất cả đánh giá",
     "DN xác nhận hoàn tất đánh giá tất cả SV",
     "POST /company-internships/submit-all-evaluations",
     "HTTP 200, trạng thái phân công chuyển sang HOAN_THANH","Pass"),
    (73,"TC073","GV chấm điểm cuối kỳ",
     "GV nhập điểm cuối kỳ cho SV",
     "POST /teacher/reports/students/:maSinhVien/grade | diem_giang_vien",
     "HTTP 200, điểm GV được lưu vào phan_cong_thuc_tap","Pass"),
    (74,"TC074","GV xem tổng hợp điểm",
     "GV xem tổng hợp điểm tất cả SV hướng dẫn",
     "GET /teacher/reports/grades-summary",
     "HTTP 200, bảng điểm tổng hợp","Pass"),

    # ─── MODULE 13: THÔNG BÁO ────────────────────────────────────────────────
    ("MODULE", "13", "THÔNG BÁO (Notifications)", "", "", "", ""),
    (75,"TC075","Lấy thông báo của tôi",
     "Người dùng lấy danh sách thông báo cá nhân",
     "GET /notifications/me | token hợp lệ",
     "HTTP 200, danh sách thông báo chưa đọc/đã đọc","Pass"),
    (76,"TC076","Đánh dấu thông báo đã đọc",
     "Người dùng đánh dấu một thông báo là đã đọc",
     "PATCH /notifications/:id | is_read: 1",
     "HTTP 200, thông báo được đánh dấu đã đọc","Pass"),
    (77,"TC077","Đánh dấu tất cả đã đọc",
     "Người dùng đánh dấu tất cả thông báo là đã đọc",
     "PATCH /notifications/bulk-update | action: mark_all_read",
     "HTTP 200, tất cả thông báo được cập nhật","Pass"),
    (78,"TC078","Xóa thông báo",
     "Người dùng xóa một thông báo khỏi danh sách",
     "DELETE /notifications/:id",
     "HTTP 200, thông báo bị xóa","Pass"),

    # ─── MODULE 14: IMPORT DỮ LIỆU ──────────────────────────────────────────
    ("MODULE", "14", "IMPORT DỮ LIỆU (Import)", "", "", "", ""),
    (79,"TC079","Tải template Excel import sinh viên",
     "Admin tải file mẫu để nhập SV hàng loạt",
     "GET /import/template",
     "HTTP 200, file Excel mẫu được tải về","Pass"),
    (80,"TC080","Validate file Excel trước khi import",
     "Admin kiểm tra dữ liệu trước khi import chính thức",
     "POST /import/validate | file Excel danh sách SV",
     "HTTP 200, danh sách lỗi (nếu có) và bản ghi hợp lệ","Pass"),
    (81,"TC081","Import sinh viên từ Excel – hợp lệ",
     "Admin import hàng loạt sinh viên từ file Excel",
     "POST /import/students | file Excel đúng định dạng",
     "HTTP 200, số bản ghi thành công / thất bại","Pass"),
    (82,"TC082","Import sinh viên – file sai định dạng",
     "Admin upload file không phải Excel",
     "POST /import/students | file PDF",
     "HTTP 400, message: 'Chỉ chấp nhận file Excel'","Pass"),

    # ─── MODULE 15: DASHBOARD & THỐNG KÊ ────────────────────────────────────
    ("MODULE", "15", "DASHBOARD & THỐNG KÊ (Dashboard)", "", "", "", ""),
    (83,"TC083","Lấy thống kê tổng quan – Admin",
     "Admin xem dashboard tổng quan toàn hệ thống",
     "GET /dashboard/stats | token admin",
     "HTTP 200, SV, GV, DN, đợt thực tập đang hoạt động","Pass"),
    (84,"TC084","Lấy thống kê công khai",
     "Xem thống kê không cần đăng nhập",
     "GET /dashboard/stats-public",
     "HTTP 200, thông tin thống kê cơ bản","Pass"),
    (85,"TC085","Thống kê sinh viên theo đợt",
     "Admin xem phân bổ sinh viên theo từng đợt thực tập",
     "GET /dashboard/students-by-period",
     "HTTP 200, biểu đồ số lượng SV theo đợt","Pass"),
    (86,"TC086","Xuất danh sách SV theo đợt",
     "Admin xuất báo cáo sinh viên theo đợt ra Excel",
     "GET /dashboard/export-students-by-period",
     "HTTP 200, file Excel được tải về","Pass"),

    # ─── MODULE 16: ZALO BOT ─────────────────────────────────────────────────
    ("MODULE", "16", "TÍCH HỢP ZALO BOT (Zalo)", "", "", "", ""),
    (87,"TC087","Nhận webhook từ Zalo",
     "Hệ thống nhận sự kiện từ Zalo OA",
     "POST /zalo/webhook | payload từ Zalo",
     "HTTP 200, sự kiện được xử lý","Pass"),
    (88,"TC088","Gửi thông báo Zalo cho sinh viên",
     "Hệ thống gửi tin nhắn Zalo thông báo deadline nộp báo cáo",
     "Hệ thống tự động trigger khi deadline trong 24h",
     "Tin nhắn Zalo được gửi tới số điện thoại SV","Pass"),

    # ─── MODULE 17: HỒ SƠ CÁ NHÂN ──────────────────────────────────────────
    ("MODULE", "17", "HỒ SƠ CÁ NHÂN (Profile)", "", "", "", ""),
    (89,"TC089","Xem hồ sơ cá nhân",
     "Người dùng xem thông tin hồ sơ của mình",
     "GET /profile/me | token hợp lệ",
     "HTTP 200, thông tin cá nhân đầy đủ","Pass"),
    (90,"TC090","Cập nhật hồ sơ cá nhân",
     "Người dùng chỉnh sửa thông tin cá nhân",
     "PUT /profile/me | so_dien_thoai, dia_chi, email_ca_nhan",
     "HTTP 200, hồ sơ được cập nhật","Pass"),
    (91,"TC091","Giảng viên xem dashboard cá nhân",
     "GV xem tổng quan SV hướng dẫn, báo cáo chờ chấm",
     "GET /teacher/profile/dashboard",
     "HTTP 200, số SV đang hướng dẫn, báo cáo chờ chấm","Pass"),
    (92,"TC092","Doanh nghiệp xem danh sách SV thực tập",
     "DN xem danh sách SV đang thực tập tại đơn vị mình",
     "GET /company-internships/my-interns",
     "HTTP 200, danh sách SV kèm thông tin thực tập","Pass"),
]

# ── Helper functions ──────────────────────────────────────────────────────────
def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(shd)

def set_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top","left","bottom","right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0"); b.set(qn("w:color"), "B0B0B0")
        tcBorders.append(b)
    tcPr.append(tcBorders)

def write_cell(cell, text, bold=False, center=False, size=9,
               color: RGBColor | None = None, italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(str(text))
    r.bold = bold; r.italic = italic
    r.font.size = Pt(size)
    r.font.name = "Times New Roman"
    if color: r.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def merge_row(table, row_idx, text, bg: RGBColor, fg: RGBColor):
    row = table.rows[row_idx]
    # merge tất cả 7 cells
    row.cells[0].merge(row.cells[6])
    cell = row.cells[0]
    set_cell_bg(cell, bg)
    set_borders(cell)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(10)
    r.font.name = "Times New Roman"
    r.font.color.rgb = fg
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# ── Main ──────────────────────────────────────────────────────────────────────
def build(output_path: str):
    doc = Document()

    # Trang A4
    sec = doc.sections[0]
    sec.page_width = Cm(29.7); sec.page_height = Cm(21)   # landscape
    sec.left_margin = Cm(2); sec.right_margin = Cm(1.5)
    sec.top_margin = Cm(2);  sec.bottom_margin = Cm(2)

    # Tiêu đề
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("BẢNG TỔNG QUAN CÁC TEST CASE HỆ THỐNG")
    r.bold = True; r.font.size = Pt(14)
    r.font.name = "Times New Roman"; r.font.color.rgb = C_HEADER

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = s.add_run("Hệ thống Quản lý Thực tập và Hợp tác Doanh nghiệp - Khoa CNTT - ĐH Đại Nam")
    sr.italic = True; sr.font.size = Pt(11); sr.font.name = "Times New Roman"

    doc.add_paragraph()

    # Caption bảng
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run("Bảng 3.1: Bảng tổng quan các test case toàn hệ thống")
    cr.bold = True; cr.font.size = Pt(11)
    cr.font.name = "Times New Roman"; cr.font.color.rgb = C_CAPTION

    # Đếm số hàng thực tế (bao gồm header + data + module separators)
    n_rows = 1 + len(TESTCASES)

    # Tạo bảng
    COL_W = [Cm(1.0), Cm(1.5), Cm(3.5), Cm(4.5), Cm(5.0), Cm(5.5), Cm(2.5)]
    HEADERS = ["STT", "Mã TC", "Chức năng / Module", "Mô tả Test Case",
               "Dữ liệu đầu vào", "Kết quả mong đợi", "Trạng thái"]

    tbl = doc.add_table(rows=n_rows, cols=7)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"

    # Header row
    hdr = tbl.rows[0].cells
    for j, (h, w) in enumerate(zip(HEADERS, COL_W)):
        hdr[j].width = w
        set_cell_bg(hdr[j], C_HEADER)
        set_borders(hdr[j])
        write_cell(hdr[j], h, bold=True, center=True, size=10, color=C_WHITE)

    # Data rows
    data_row = 1
    odd = True
    for tc in TESTCASES:
        row = tbl.rows[data_row]

        if tc[0] == "MODULE":
            # Module separator
            merge_row(tbl, data_row,
                      f"MODULE {tc[1]}: {tc[2]}",
                      C_MODULE_BG, C_MODULE)
            odd = True
        else:
            stt, ma_tc, chuc_nang, mo_ta, du_lieu, ket_qua, trang_thai = tc
            bg = C_ODD if odd else None
            vals = [stt, ma_tc, chuc_nang, mo_ta, du_lieu, ket_qua, trang_thai]
            for j, (val, w) in enumerate(zip(vals, COL_W)):
                c = row.cells[j]
                c.width = w
                if bg: set_cell_bg(c, bg)
                set_borders(c)
                is_center = j in (0, 1, 6)
                fc = C_PASS if (j == 6 and val == "Pass") else None
                write_cell(c, val, center=is_center, size=9,
                           color=fc, bold=(j == 6 and val == "Pass"))
            odd = not odd

        data_row += 1

    doc.save(output_path)
    print(f"[OK] Da xuat: {output_path}")

if __name__ == "__main__":
    out = os.path.join(
        os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
        "backend", "docs", "Bang_Tong_Quan_TestCase.docx"
    )
    build(out)
