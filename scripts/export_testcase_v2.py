# -*- coding: utf-8 -*-
"""
Xuất bảng test case chi tiết (format mới) theo hình 2:
STT | Mục tiêu kiểm thử | Mô tả thao tác | Kết quả chờ đợi |
Dữ liệu kiểm thử | Kết quả thực tế | [Kiểm thử lần 1: P/F/N | Bug# | Diễn giải]
"""

import os
import sys

# Đảm bảo stdout UTF-8 trên Windows
if sys.stdout.encoding != 'utf-8':
    try:
        sys.stdout.reconfigure(encoding='utf-8')
    except Exception:
        pass

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Màu sắc ──────────────────────────────────────────────────────────────────
C_HEADER    = RGBColor(0x1F, 0x49, 0x7D)   # xanh đậm header
C_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
C_MODULE_BG = RGBColor(0xBD, 0xD7, 0xEE)   # nền section module
C_MODULE_FG = RGBColor(0x1F, 0x49, 0x7D)   # chữ section module
C_ODD       = RGBColor(0xED, 0xF3, 0xFB)   # hàng lẻ
C_PASS      = RGBColor(0x00, 0x70, 0x00)   # màu Pass
C_CAPTION   = RGBColor(0x1F, 0x49, 0x7D)

# ── Dữ liệu test case ────────────────────────────────────────────────────────
# (stt, tên_tc, mô_tả_thao_tác, kết_quả_chờ_đợi, dữ_liệu_kiểm_thử, kết_quả_thực_tế, p/f/n)
# MODULE row: ("MODULE", số_module, tên_module)
TESTCASES = [

    # ─── MODULE 1: XÁC THỰC HỆ THỐNG (9 TC) ────────────────────────────────
    ("MODULE", "1", "Chức năng Xác thực hệ thống (Đăng nhập / Đăng xuất)"),
    (1,
     "Đăng nhập thành công với tài khoản sinh viên",
     "1. Mở trình duyệt, truy cập trang đăng nhập.\n2. Nhập userId: SV001.\n3. Nhập password: 123456.\n4. Chọn vai trò: Sinh viên.\n5. Nhấn nút Đăng nhập.",
     "Hệ thống trả về token JWT, thông tin sinh viên, HTTP 200. Chuyển đến trang chủ sinh viên",
     "userId: SV001 | password: 123456 | role: sinh-vien",
     "Hệ thống trả về token JWT và thông tin sinh viên thành công, HTTP 200",
     "P"),
    (2,
     "Đăng nhập thành công với tài khoản quản trị viên",
     "1. Mở trang đăng nhập Admin.\n2. Nhập userId: admin001.\n3. Nhập password: 123456.\n4. Chọn vai trò: Admin.\n5. Nhấn nút Đăng nhập.",
     "Hệ thống trả về token JWT, thông tin admin, HTTP 200. Chuyển đến trang Dashboard admin",
     "userId: admin001 | password: 123456 | role: admin",
     "Hệ thống trả về token JWT và thông tin admin thành công, HTTP 200",
     "P"),
    (3,
     "Đăng nhập thành công với tài khoản giảng viên",
     "1. Mở trang đăng nhập.\n2. Nhập userId: GV001.\n3. Nhập password: 123456.\n4. Chọn vai trò: Giảng viên.\n5. Nhấn nút Đăng nhập.",
     "Hệ thống trả về token JWT, thông tin giảng viên, HTTP 200",
     "userId: GV001 | password: 123456 | role: giang-vien",
     "Hệ thống trả về token JWT và thông tin giảng viên thành công, HTTP 200",
     "P"),
    (4,
     "Đăng nhập thành công với tài khoản doanh nghiệp",
     "1. Mở trang đăng nhập.\n2. Nhập userId: DN001.\n3. Nhập password: 123456.\n4. Chọn vai trò: Doanh nghiệp.\n5. Nhấn nút Đăng nhập.",
     "Hệ thống trả về token JWT, thông tin doanh nghiệp, HTTP 200",
     "userId: DN001 | password: 123456 | role: doanh-nghiep",
     "Hệ thống trả về token JWT và thông tin doanh nghiệp thành công, HTTP 200",
     "P"),
    (5,
     "Đăng nhập thất bại vì sai mật khẩu",
     "1. Mở trang đăng nhập.\n2. Nhập userId: SV001.\n3. Nhập password sai: sai123.\n4. Chọn vai trò: Sinh viên.\n5. Nhấn nút Đăng nhập.",
     "Hệ thống trả về lỗi HTTP 401, hiển thị thông báo 'Mật khẩu không chính xác'",
     "userId: SV001 | password: sai123 | role: sinh-vien",
     "HTTP 401, hiển thị thông báo 'Mật khẩu không chính xác'",
     "P"),
    (6,
     "Đăng nhập thất bại vì tài khoản bị khóa",
     "1. Mở trang đăng nhập.\n2. Nhập userId của tài khoản bị khóa: SV_LOCKED.\n3. Nhập password: 123456.\n4. Nhấn nút Đăng nhập.\n5. Quan sát thông báo trả về.",
     "Hệ thống trả về lỗi HTTP 403, thông báo 'Tài khoản đã bị vô hiệu hóa'",
     "userId: SV_LOCKED | password: 123456 | is_active: 0",
     "HTTP 403, hiển thị thông báo 'Tài khoản đã bị vô hiệu hóa'",
     "P"),
    (7,
     "Đăng nhập thất bại khi không nhập thông tin",
     "1. Mở trang đăng nhập.\n2. Để trống trường userId.\n3. Để trống trường password.\n4. Nhấn nút Đăng nhập.\n5. Quan sát thông báo lỗi validation.",
     "Hệ thống hiển thị thông báo yêu cầu nhập đầy đủ thông tin, không gửi request",
     "userId: (trống) | password: (trống)",
     "Hiển thị thông báo validation, không gửi request lên server",
     "P"),
    (8,
     "Đăng xuất khỏi hệ thống",
     "1. Đăng nhập vào hệ thống thành công.\n2. Nhấn vào avatar/tên người dùng góc trên phải.\n3. Chọn mục Đăng xuất.\n4. Xác nhận đăng xuất (nếu có).\n5. Quan sát kết quả.",
     "Hệ thống hủy token, chuyển về trang đăng nhập, HTTP 200",
     "Token JWT hợp lệ trong Authorization header",
     "HTTP 200, token bị hủy, chuyển về trang đăng nhập",
     "P"),
    (9,
     "Truy cập API bảo mật không có token",
     "1. Mở công cụ kiểm thử API (Postman/Thunder Client).\n2. Tạo request GET tới /api/profile/me.\n3. Không thêm Authorization header.\n4. Gửi request.\n5. Quan sát mã HTTP và thông báo trả về.",
     "Hệ thống trả về HTTP 401, thông báo 'Unauthorized'",
     "Không có Authorization header",
     "HTTP 401, message: 'Unauthorized'",
     "P"),

    # ─── MODULE 2: QUẢN LÝ TÀI KHOẢN VÀ MẬT KHẨU (9 TC) ───────────────────
    ("MODULE", "2", "Chức năng Quản lý tài khoản và mật khẩu"),
    (10,
     "Admin tạo tài khoản sinh viên mới",
     "1. Đăng nhập với quyền Admin.\n2. Vào menu Quản lý tài khoản.\n3. Nhấn nút Thêm tài khoản.\n4. Nhập đầy đủ thông tin: userId, email, password, role.\n5. Nhấn Lưu.",
     "Tài khoản được tạo thành công, HTTP 201, dữ liệu lưu vào CSDL",
     "userId: SV100 | email: sv100@dainam.edu.vn | password: 123456 | role: sinh-vien",
     "HTTP 201, tài khoản được tạo, dữ liệu lưu vào bảng accounts",
     "P"),
    (11,
     "Tạo tài khoản thất bại vì userId đã tồn tại",
     "1. Đăng nhập với quyền Admin.\n2. Vào menu Quản lý tài khoản.\n3. Nhấn nút Thêm tài khoản.\n4. Nhập userId đã tồn tại trong hệ thống: SV001.\n5. Nhấn Lưu và quan sát thông báo.",
     "Hệ thống trả về HTTP 400, thông báo 'userId đã tồn tại'",
     "userId: SV001 (đã tồn tại trong CSDL)",
     "HTTP 400, message: 'userId đã tồn tại'",
     "P"),
    (12,
     "Admin lấy danh sách tài khoản có phân trang",
     "1. Đăng nhập với quyền Admin.\n2. Vào menu Quản lý tài khoản.\n3. Quan sát danh sách tài khoản hiển thị.\n4. Kiểm tra phân trang (page, limit).\n5. Chuyển sang trang tiếp theo.",
     "Hệ thống trả về danh sách tài khoản có phân trang, HTTP 200",
     "GET /api/accounts?page=1&limit=10 | token admin",
     "HTTP 200, danh sách tài khoản có thông tin phân trang",
     "P"),
    (13,
     "Admin kích hoạt lại tài khoản bị khóa",
     "1. Đăng nhập với quyền Admin.\n2. Vào danh sách tài khoản, tìm tài khoản bị khóa.\n3. Nhấn nút Kích hoạt bên cạnh tài khoản.\n4. Xác nhận thao tác.\n5. Kiểm tra trạng thái tài khoản.",
     "Tài khoản được kích hoạt, is_active = 1, HTTP 200",
     "POST /api/accounts/:userId/activate | token admin",
     "HTTP 200, is_active = 1, tài khoản có thể đăng nhập",
     "P"),
    (14,
     "Admin vô hiệu hóa tài khoản",
     "1. Đăng nhập với quyền Admin.\n2. Vào danh sách tài khoản, tìm tài khoản cần khóa.\n3. Nhấn nút Vô hiệu hóa bên cạnh tài khoản.\n4. Xác nhận thao tác.\n5. Kiểm tra trạng thái tài khoản.",
     "Tài khoản bị khóa, is_active = 0, HTTP 200",
     "POST /api/accounts/:userId/deactivate | token admin",
     "HTTP 200, is_active = 0, tài khoản không thể đăng nhập",
     "P"),
    (15,
     "Người dùng đổi mật khẩu thành công",
     "1. Đăng nhập vào hệ thống.\n2. Vào trang Hồ sơ cá nhân > Bảo mật.\n3. Nhấn nút Đổi mật khẩu.\n4. Nhập mật khẩu cũ: 123456.\n5. Nhập mật khẩu mới: NewPass@123.\n6. Xác nhận mật khẩu mới.\n7. Nhấn Lưu.",
     "Mật khẩu được cập nhật, HTTP 200",
     "oldPassword: 123456 | newPassword: NewPass@123 | confirmPassword: NewPass@123",
     "HTTP 200, mật khẩu được cập nhật, đăng nhập lại với mật khẩu mới thành công",
     "P"),
    (16,
     "Gửi mã xác minh quên mật khẩu qua email",
     "1. Mở trang đăng nhập, nhấn Quên mật khẩu.\n2. Nhập địa chỉ email đã đăng ký.\n3. Nhấn nút Gửi mã xác minh.\n4. Kiểm tra hộp thư email.\n5. Xác nhận email chứa mã 6 chữ số.",
     "Hệ thống gửi mã xác nhận 6 chữ số về email, HTTP 200",
     "email: sv001@dainam.edu.vn",
     "HTTP 200, email chứa mã xác nhận được gửi đến hộp thư",
     "P"),
    (17,
     "Xác minh mã reset mật khẩu đúng",
     "1. Sau bước gửi mã, mở email nhận được.\n2. Sao chép mã xác nhận 6 chữ số.\n3. Quay lại trang đặt lại mật khẩu.\n4. Nhập mã xác nhận vào ô.\n5. Nhấn Xác minh.",
     "Mã hợp lệ, hệ thống cho phép đặt mật khẩu mới, HTTP 200",
     "email: sv001@dainam.edu.vn | code: 123456 (đúng)",
     "HTTP 200, xác nhận thành công, cho phép nhập mật khẩu mới",
     "P"),
    (18,
     "Đặt lại mật khẩu mới sau xác minh",
     "1. Sau khi xác minh mã thành công.\n2. Nhập mật khẩu mới: NewPass@2024.\n3. Nhập lại mật khẩu xác nhận: NewPass@2024.\n4. Nhấn nút Đặt lại mật khẩu.\n5. Đăng nhập lại với mật khẩu mới để kiểm tra.",
     "Mật khẩu được đặt lại thành công, HTTP 200",
     "email, code hợp lệ | newPassword: NewPass@2024 | confirmPassword: NewPass@2024",
     "HTTP 200, mật khẩu được cập nhật, đăng nhập với mật khẩu mới thành công",
     "P"),

    # ─── MODULE 3: QUẢN LÝ SINH VIÊN (9 TC) ────────────────────────────────
    ("MODULE", "3", "Chức năng Quản lý sinh viên"),
    (19,
     "Admin xem danh sách sinh viên có phân trang",
     "1. Đăng nhập với quyền Admin.\n2. Vào menu Quản lý sinh viên.\n3. Quan sát danh sách hiển thị.\n4. Kiểm tra thông tin phân trang.\n5. Chuyển trang và lọc theo khoa.",
     "Danh sách sinh viên được hiển thị có phân trang, HTTP 200",
     "GET /api/sinh-vien?page=1&limit=20 | token admin",
     "HTTP 200, danh sách sinh viên với thông tin cơ bản",
     "P"),
    (20,
     "Xem chi tiết thông tin sinh viên theo ID",
     "1. Đăng nhập với quyền Admin.\n2. Vào danh sách sinh viên.\n3. Tìm sinh viên cần xem.\n4. Nhấn vào tên hoặc nút Chi tiết.\n5. Quan sát thông tin đầy đủ.",
     "Thông tin đầy đủ của sinh viên được hiển thị, HTTP 200",
     "GET /api/sinh-vien/:id | token admin",
     "HTTP 200, trả về đầy đủ thông tin sinh viên",
     "P"),
    (21,
     "Admin thêm sinh viên mới vào hệ thống",
     "1. Đăng nhập với quyền Admin.\n2. Vào menu Quản lý sinh viên.\n3. Nhấn nút Thêm sinh viên.\n4. Nhập đầy đủ thông tin bắt buộc.\n5. Nhấn Lưu và kiểm tra kết quả.",
     "Sinh viên được tạo thành công, HTTP 201",
     "ma_sinh_vien: SV100 | ho_ten: Nguyễn Văn A | lop: CNTT1 | khoa: CNTT | account_id hợp lệ",
     "HTTP 201, sinh viên được lưu vào CSDL",
     "P"),
    (22,
     "Cập nhật thông tin sinh viên",
     "1. Đăng nhập (SV hoặc Admin).\n2. Vào trang hồ sơ sinh viên.\n3. Nhấn nút Chỉnh sửa.\n4. Sửa thông tin: số điện thoại, địa chỉ.\n5. Nhấn Lưu và kiểm tra dữ liệu.",
     "Thông tin được cập nhật, HTTP 200",
     "PUT /api/sinh-vien/:id | so_dien_thoai, dia_chi cập nhật",
     "HTTP 200, thông tin sinh viên được cập nhật trong CSDL",
     "P"),
    (23,
     "Admin xóa sinh viên khỏi hệ thống",
     "1. Đăng nhập với quyền Admin.\n2. Vào danh sách sinh viên.\n3. Tìm sinh viên cần xóa.\n4. Nhấn nút Xóa bên cạnh sinh viên.\n5. Xác nhận xóa trong hộp thoại.",
     "Sinh viên bị xóa, HTTP 200",
     "DELETE /api/sinh-vien/:id | token admin",
     "HTTP 200, sinh viên bị xóa khỏi CSDL",
     "P"),
    (24,
     "Sinh viên tải lên CV định dạng PDF hợp lệ",
     "1. Đăng nhập với tài khoản sinh viên.\n2. Vào trang Đăng ký thực tập.\n3. Nhấn nút Tải CV lên.\n4. Chọn file CV định dạng PDF, kích thước < 5MB.\n5. Nhấn Upload và kiểm tra kết quả.",
     "CV được lưu lên server, đường dẫn lưu vào CSDL, HTTP 200",
     "Multipart form-data | file: CV_SV001.pdf | kích thước: 2MB",
     "HTTP 200, cv_path được lưu vào CSDL, file lưu tại /uploads/cv/",
     "P"),
    (25,
     "Từ chối upload CV sai định dạng",
     "1. Đăng nhập với tài khoản sinh viên.\n2. Vào trang Đăng ký thực tập.\n3. Nhấn nút Tải CV lên.\n4. Chọn file không phải PDF (ví dụ: picture.jpg).\n5. Nhấn Upload và quan sát thông báo lỗi.",
     "Hệ thống từ chối, hiển thị thông báo lỗi, HTTP 400",
     "Multipart form-data | file: picture.jpg",
     "HTTP 400, message: 'Định dạng file không hợp lệ'",
     "P"),
    (26,
     "Phân tích CV bằng AI tự động",
     "1. Đăng nhập với tài khoản sinh viên.\n2. Đảm bảo đã upload CV thành công.\n3. Nhấn nút Phân tích CV.\n4. Chờ hệ thống xử lý.\n5. Quan sát kết quả phân tích kỹ năng.",
     "Kết quả phân tích kỹ năng và thông tin từ CV được hiển thị, HTTP 200",
     "POST /api/sinh-vien/analyze-cv | cv_path hợp lệ trong CSDL",
     "HTTP 200, trả về danh sách kỹ năng trích xuất từ CV",
     "P"),
    (27,
     "Lấy thống kê tổng quan sinh viên",
     "1. Đăng nhập với quyền Admin.\n2. Vào trang Dashboard.\n3. Quan sát khung thống kê sinh viên.\n4. Kiểm tra số liệu tổng, đã phân công, chưa phân công.\n5. So sánh với dữ liệu thực tế trong CSDL.",
     "HTTP 200, trả về tổng số SV, số đã phân công, chưa phân công",
     "GET /api/sinh-vien/stats | token admin",
     "HTTP 200, tổng số SV, đã phân công, chưa phân công chính xác",
     "P"),

    # ─── MODULE 4: QUẢN LÝ GIẢNG VIÊN (5 TC) ───────────────────────────────
    ("MODULE", "4", "Chức năng Quản lý giảng viên"),
    (28,
     "Admin lấy danh sách giảng viên",
     "1. Đăng nhập với quyền Admin.\n2. Vào menu Quản lý giảng viên.\n3. Quan sát danh sách hiển thị.\n4. Thử tìm kiếm theo tên giảng viên.\n5. Kiểm tra thông tin phân trang.",
     "Danh sách giảng viên hiển thị đầy đủ, HTTP 200",
     "GET /api/giang-vien?page=1&limit=20 | token admin",
     "HTTP 200, danh sách giảng viên với thông tin đầy đủ",
     "P"),
    (29,
     "Admin thêm giảng viên mới",
     "1. Đăng nhập với quyền Admin.\n2. Vào menu Quản lý giảng viên.\n3. Nhấn nút Thêm giảng viên.\n4. Nhập đầy đủ thông tin: mã GV, họ tên, khoa, học vị.\n5. Nhấn Lưu và kiểm tra kết quả.",
     "Giảng viên được thêm vào hệ thống, HTTP 201",
     "ma_giang_vien: GV010 | ho_ten: PGS.TS Trần Văn B | khoa: CNTT | hoc_vi: Tiến sĩ",
     "HTTP 201, giảng viên được lưu vào CSDL",
     "P"),
    (30,
     "Cập nhật thông tin giảng viên",
     "1. Đăng nhập (Admin hoặc Giảng viên).\n2. Vào trang hồ sơ giảng viên.\n3. Nhấn nút Chỉnh sửa.\n4. Sửa bộ môn, chuyên môn, số điện thoại.\n5. Nhấn Lưu và kiểm tra dữ liệu.",
     "Thông tin giảng viên được cập nhật, HTTP 200",
     "PUT /api/giang-vien/:id | bo_mon, chuyen_mon, so_dien_thoai",
     "HTTP 200, thông tin giảng viên cập nhật thành công",
     "P"),
    (31,
     "Tự động phân công giảng viên cho sinh viên",
     "1. Đăng nhập với quyền Admin.\n2. Vào trang Phân công thực tập.\n3. Nhấn nút Tự động phân công giảng viên.\n4. Xác nhận thao tác.\n5. Kiểm tra số lượng SV được phân công và tải của từng GV.",
     "Hệ thống phân GV dựa trên tải (TS: tối đa 20 SV, ThS: 15 SV), HTTP 200",
     "POST /api/giang-vien/auto-assign | token admin",
     "HTTP 200, số lượng SV được phân công GV, GV không quá tải",
     "P"),
    (32,
     "Xuất danh sách giảng viên ra Excel",
     "1. Đăng nhập với quyền Admin.\n2. Vào trang Quản lý giảng viên.\n3. Nhấn nút Xuất Excel.\n4. Chờ file được tạo.\n5. Kiểm tra file Excel tải về có đủ dữ liệu.",
     "File Excel chứa danh sách giảng viên được tải về, HTTP 200",
     "GET /api/giang-vien/export | token admin",
     "HTTP 200, file Excel hợp lệ được tải về máy tính",
     "P"),

    # ─── MODULE 5: QUẢN LÝ DOANH NGHIỆP (5 TC) ─────────────────────────────
    ("MODULE", "5", "Chức năng Quản lý doanh nghiệp"),
    (33,
     "Admin xem danh sách doanh nghiệp",
     "1. Đăng nhập với quyền Admin.\n2. Vào menu Quản lý doanh nghiệp.\n3. Quan sát danh sách hiển thị.\n4. Thử tìm kiếm theo tên công ty.\n5. Kiểm tra thông tin phân trang.",
     "Danh sách doanh nghiệp hiển thị có phân trang, HTTP 200",
     "GET /api/doanh-nghiep?page=1&limit=20 | token admin",
     "HTTP 200, danh sách doanh nghiệp với thông tin đầy đủ",
     "P"),
    (34,
     "Admin thêm doanh nghiệp hợp tác mới",
     "1. Đăng nhập với quyền Admin.\n2. Vào menu Quản lý doanh nghiệp.\n3. Nhấn nút Thêm doanh nghiệp.\n4. Nhập tên công ty, địa chỉ, số điện thoại.\n5. Nhấn Lưu và kiểm tra kết quả.",
     "Doanh nghiệp được thêm thành công, HTTP 201",
     "ten_cong_ty: Công ty ABC | dia_chi: Hà Nội | so_dien_thoai: 0241234567",
     "HTTP 201, doanh nghiệp được lưu vào CSDL",
     "P"),
    (35,
     "Cập nhật thông tin doanh nghiệp",
     "1. Đăng nhập (Admin hoặc Doanh nghiệp).\n2. Vào trang hồ sơ doanh nghiệp.\n3. Nhấn nút Chỉnh sửa.\n4. Sửa địa chỉ, số điện thoại, website.\n5. Nhấn Lưu và kiểm tra dữ liệu.",
     "Thông tin doanh nghiệp được cập nhật, HTTP 200",
     "PUT /api/doanh-nghiep/:id | dia_chi, so_dien_thoai, website",
     "HTTP 200, thông tin doanh nghiệp cập nhật thành công",
     "P"),
    (36,
     "Admin xóa doanh nghiệp khỏi hệ thống",
     "1. Đăng nhập với quyền Admin.\n2. Vào danh sách doanh nghiệp.\n3. Tìm doanh nghiệp cần xóa.\n4. Nhấn nút Xóa bên cạnh doanh nghiệp.\n5. Xác nhận xóa trong hộp thoại.",
     "Doanh nghiệp bị xóa, HTTP 200",
     "DELETE /api/doanh-nghiep/:id | token admin",
     "HTTP 200, doanh nghiệp bị xóa khỏi CSDL",
     "P"),
    (37,
     "Xuất danh sách doanh nghiệp ra Excel",
     "1. Đăng nhập với quyền Admin.\n2. Vào trang Quản lý doanh nghiệp.\n3. Nhấn nút Xuất Excel.\n4. Chờ file được tạo.\n5. Kiểm tra file Excel tải về có đủ dữ liệu.",
     "File Excel chứa danh sách doanh nghiệp được tải về, HTTP 200",
     "GET /api/doanh-nghiep/export | token admin",
     "HTTP 200, file Excel hợp lệ được tải về máy tính",
     "P"),

    # ─── MODULE 6: ĐĂNG KÝ THỰC TẬP (6 TC) ─────────────────────────────────
    ("MODULE", "6", "Chức năng Đăng ký thực tập"),
    (38,
     "Sinh viên đăng ký thực tập theo kiểu khoa giới thiệu",
     "1. Đăng nhập với tài khoản sinh viên.\n2. Vào menu Đăng ký thực tập.\n3. Chọn nguyện vọng: Khoa giới thiệu.\n4. Nhập vị trí mong muốn, đính kèm CV.\n5. Nhấn Nộp đơn và kiểm tra trạng thái.",
     "Đơn đăng ký được tạo, trạng thái: CHỜ_DUYỆT, HTTP 201",
     "nguyen_vong: khoa-gioi-thieu | vi_tri_muon_ung_tuyen: Lập trình viên | file CV đính kèm",
     "HTTP 201, đơn đăng ký được tạo, trạng thái CHỜ_DUYỆT",
     "P"),
    (39,
     "Sinh viên đăng ký thực tập tự liên hệ công ty",
     "1. Đăng nhập với tài khoản sinh viên.\n2. Vào menu Đăng ký thực tập.\n3. Chọn nguyện vọng: Tự liên hệ.\n4. Nhập tên công ty, địa chỉ, người liên hệ.\n5. Nhấn Nộp đơn và kiểm tra trạng thái.",
     "Đơn đăng ký được tạo với thông tin công ty, trạng thái: CHỜ_DUYỆT, HTTP 201",
     "nguyen_vong: tu-lien-he | ten_cong_ty: Công ty XYZ | dia_chi_cong_ty: Hà Nội | nguoi_lien_he",
     "HTTP 201, đơn đăng ký với thông tin công ty tự liên hệ được lưu",
     "P"),
    (40,
     "Sinh viên kiểm tra trạng thái đơn đăng ký",
     "1. Đăng nhập với tài khoản sinh viên.\n2. Vào trang Hồ sơ cá nhân.\n3. Kéo xuống phần Thông tin đăng ký thực tập.\n4. Quan sát trạng thái đơn đăng ký.\n5. Kiểm tra thông tin chi tiết đơn.",
     "Hiển thị trạng thái và thông tin đơn đăng ký, HTTP 200",
     "GET /api/registration/check | token sinh-vien",
     "HTTP 200, trả về trạng thái và đầy đủ thông tin đơn đăng ký",
     "P"),
    (41,
     "Admin duyệt đơn đăng ký thực tập",
     "1. Đăng nhập với quyền Admin.\n2. Vào menu Quản lý đăng ký thực tập.\n3. Tìm đơn đăng ký có trạng thái CHỜ_DUYỆT.\n4. Nhấn nút Duyệt bên cạnh đơn.\n5. Xác nhận và kiểm tra trạng thái mới.",
     "Đơn được duyệt, trạng thái chuyển sang ĐÃ_DUYỆT, HTTP 200",
     "PUT /api/admin/student-registrations/:id/approve | token admin",
     "HTTP 200, trạng thái: ĐÃ_DUYỆT, sinh viên nhận thông báo",
     "P"),
    (42,
     "Admin từ chối đơn đăng ký với lý do cụ thể",
     "1. Đăng nhập với quyền Admin.\n2. Vào danh sách đơn đăng ký.\n3. Tìm đơn cần từ chối.\n4. Nhấn nút Từ chối, nhập lý do: 'CV chưa hợp lệ'.\n5. Xác nhận và kiểm tra trạng thái mới.",
     "Đơn bị từ chối, trạng thái: TỪ_CHỐI, lý do được lưu, HTTP 200",
     "PUT /api/admin/student-registrations/:id/reject | ly_do_tu_choi: 'CV chưa hợp lệ'",
     "HTTP 200, trạng thái: TỪ_CHỐI, lý do từ chối được lưu vào CSDL",
     "P"),
    (43,
     "Admin duyệt hàng loạt theo nguyện vọng",
     "1. Đăng nhập với quyền Admin.\n2. Vào trang Quản lý đăng ký thực tập.\n3. Nhấn nút Duyệt hàng loạt.\n4. Chọn lọc theo nguyện vọng: Khoa giới thiệu.\n5. Xác nhận và kiểm tra số lượng đơn được duyệt.",
     "Số lượng SV được duyệt trả về, HTTP 200",
     "POST /api/admin/students/bulk-approve-by-preference | nguyen_vong: khoa-gioi-thieu",
     "HTTP 200, số lượng đơn đã duyệt chính xác",
     "P"),

    # ─── MODULE 7: QUY TRÌNH PHỎNG VẤN (5 TC) ──────────────────────────────
    ("MODULE", "7", "Chức năng Quy trình phỏng vấn (Interview Workflow)"),
    (44,
     "Sinh viên nộp đơn ứng tuyển vào doanh nghiệp",
     "1. Đăng nhập với tài khoản sinh viên.\n2. Vào trang Phỏng vấn / Ứng tuyển.\n3. Chọn tin tuyển dụng phù hợp.\n4. Viết thư xin việc và đính kèm CV.\n5. Nhấn Ứng tuyển và kiểm tra trạng thái.",
     "Đơn ứng tuyển được tạo, trạng thái: ĐANG_CHỜ, HTTP 201",
     "POST /api/interview-workflow/student/applications | tin_tuyen_dung_id, thu_xin_viec",
     "HTTP 201, đơn ứng tuyển được lưu, trạng thái PENDING",
     "P"),
    (45,
     "Sinh viên xem danh sách đơn ứng tuyển của mình",
     "1. Đăng nhập với tài khoản sinh viên.\n2. Vào trang Phỏng vấn.\n3. Chọn tab Đơn ứng tuyển của tôi.\n4. Quan sát danh sách và trạng thái từng đơn.\n5. Nhấn vào đơn để xem chi tiết.",
     "Danh sách đơn ứng tuyển với trạng thái được hiển thị, HTTP 200",
     "GET /api/interview-workflow/student/applications | token sinh-vien",
     "HTTP 200, danh sách đơn ứng tuyển với trạng thái hiện tại",
     "P"),
    (46,
     "Admin duyệt hồ sơ ứng tuyển của sinh viên",
     "1. Đăng nhập với quyền Admin.\n2. Vào trang Quản lý ứng tuyển.\n3. Tìm đơn có trạng thái PENDING.\n4. Xem hồ sơ SV, nhấn nút Duyệt.\n5. Xác nhận và kiểm tra trạng thái mới.",
     "Trạng thái chuyển sang APPROVED, SV và DN được thông báo, HTTP 200",
     "PUT /api/interview-workflow/admin/applications/:id/review | action: APPROVED",
     "HTTP 200, trạng thái: APPROVED, thông báo được gửi",
     "P"),
    (47,
     "Doanh nghiệp lên lịch phỏng vấn cho sinh viên",
     "1. Đăng nhập với tài khoản doanh nghiệp.\n2. Vào trang Quản lý ứng tuyển.\n3. Tìm đơn có trạng thái APPROVED.\n4. Nhấn nút Lên lịch phỏng vấn.\n5. Nhập ngày, giờ, địa điểm rồi nhấn Xác nhận.",
     "Lịch phỏng vấn được lưu, trạng thái: ĐÃ_LÊN_LỊCH, HTTP 200",
     "PUT /api/interview-workflow/company/applications/:id/interview | ngay: 2024-12-15, gio: 14:00, dia_diem: Văn phòng tầng 3",
     "HTTP 200, lịch phỏng vấn lưu thành công, trạng thái: INTERVIEW_SCHEDULED",
     "P"),
    (48,
     "Doanh nghiệp cập nhật kết quả phỏng vấn",
     "1. Đăng nhập với tài khoản doanh nghiệp.\n2. Vào trang Quản lý ứng tuyển.\n3. Tìm đơn có trạng thái INTERVIEW_SCHEDULED.\n4. Nhấn nút Cập nhật kết quả.\n5. Chọn kết quả: Đạt (PASS) và nhấn Lưu.",
     "Kết quả được lưu, trạng thái chuyển sang PASS hoặc FAIL, HTTP 200",
     "PUT /api/interview-workflow/company/applications/:id/result | result: PASS",
     "HTTP 200, kết quả phỏng vấn PASS được lưu vào CSDL",
     "P"),

    # ─── MODULE 8: PHÂN CÔNG THỰC TẬP (4 TC) ───────────────────────────────
    ("MODULE", "8", "Chức năng Phân công thực tập"),
    (49,
     "Admin tạo phân công thực tập thủ công",
     "1. Đăng nhập với quyền Admin.\n2. Vào trang Phân công thực tập.\n3. Nhấn nút Tạo phân công.\n4. Chọn sinh viên, doanh nghiệp, giảng viên hướng dẫn.\n5. Nhập ngày bắt đầu, ngày kết thúc rồi nhấn Lưu.",
     "Phân công được tạo thành công, HTTP 201",
     "sinh_vien_id | doanh_nghiep_id | giang_vien_id | dot_thuc_tap_id | ngay_bat_dau: 2024-12-01 | ngay_ket_thuc: 2025-02-28",
     "HTTP 201, phân công được lưu vào bảng phan_cong_thuc_tap",
     "P"),
    (50,
     "Xem danh sách phân công thực tập",
     "1. Đăng nhập (Admin hoặc Giảng viên).\n2. Vào trang Phân công thực tập.\n3. Quan sát danh sách phân công.\n4. Thử lọc theo đợt, khoa, giảng viên.\n5. Kiểm tra thông tin SV-DN-GV trong mỗi phân công.",
     "Danh sách phân công với thông tin SV-DN-GV được hiển thị, HTTP 200",
     "GET /api/assignments?page=1&limit=20 | token hợp lệ",
     "HTTP 200, danh sách phân công có đầy đủ thông tin",
     "P"),
    (51,
     "Tự động phân công công ty theo kiểu khoa giới thiệu",
     "1. Đăng nhập với quyền Admin.\n2. Vào trang Phân công thực tập.\n3. Nhấn nút Tự động phân công công ty.\n4. Xác nhận thao tác trong hộp thoại.\n5. Kiểm tra số SV được phân công DN.",
     "Hệ thống tự động phân công, trả về số lượng đã phân, HTTP 200",
     "POST /api/admin/students/auto-assign-company-khoa-gioi-thieu | token admin",
     "HTTP 200, số SV được phân công DN chính xác",
     "P"),
    (52,
     "Admin hủy phân công thực tập",
     "1. Đăng nhập với quyền Admin.\n2. Vào danh sách phân công thực tập.\n3. Tìm phân công cần hủy.\n4. Nhấn nút Xóa bên cạnh phân công.\n5. Xác nhận xóa trong hộp thoại.",
     "Phân công bị xóa, HTTP 200",
     "DELETE /api/assignments/:id | token admin",
     "HTTP 200, phân công bị xóa khỏi CSDL",
     "P"),

    # ─── MODULE 9: QUẢN LÝ ĐỢT THỰC TẬP (6 TC) ─────────────────────────────
    ("MODULE", "9", "Chức năng Quản lý đợt thực tập"),
    (53,
     "Admin tạo đợt thực tập mới",
     "1. Đăng nhập với quyền Admin.\n2. Vào menu Quản lý đợt thực tập.\n3. Nhấn nút Tạo đợt mới.\n4. Nhập tên đợt, thời gian bắt đầu, kết thúc, mô tả.\n5. Nhấn Lưu và kiểm tra trạng thái.",
     "Đợt thực tập được tạo, trạng thái: sắp mở, HTTP 201",
     "ten_dot: Đợt 1 - Năm 2024-2025 | thoi_gian_bat_dau: 2024-12-01 | thoi_gian_ket_thuc: 2025-02-28",
     "HTTP 201, đợt thực tập được tạo với trạng thái sắp mở",
     "P"),
    (54,
     "Lấy danh sách tất cả đợt thực tập",
     "1. Đăng nhập vào hệ thống.\n2. Vào trang Đợt thực tập.\n3. Quan sát danh sách hiển thị.\n4. Kiểm tra thông tin từng đợt (tên, thời gian, trạng thái).\n5. Thử lọc theo trạng thái.",
     "Danh sách các đợt thực tập được hiển thị, HTTP 200",
     "GET /api/internship-batches | token hợp lệ",
     "HTTP 200, danh sách các đợt thực tập với trạng thái",
     "P"),
    (55,
     "Lấy thông tin đợt thực tập đang hoạt động",
     "1. Đăng nhập với tài khoản sinh viên.\n2. Vào trang Đăng ký thực tập.\n3. Hệ thống tự động hiển thị đợt đang mở.\n4. Kiểm tra thông tin đợt: tên, thời gian, mô tả.\n5. Xác nhận trạng thái là đang mở.",
     "Trả về thông tin đợt đang mở (trạng thái: đang mở), HTTP 200",
     "GET /api/internship-batches/active | token hợp lệ",
     "HTTP 200, trả về đợt thực tập có trạng thái đang mở",
     "P"),
    (56,
     "Admin import danh sách tham gia đợt từ Excel",
     "1. Đăng nhập với quyền Admin.\n2. Vào chi tiết đợt thực tập.\n3. Nhấn nút Import danh sách tham gia.\n4. Chọn file Excel đúng định dạng.\n5. Nhấn Import và kiểm tra số bản ghi.",
     "Dữ liệu được import thành công, HTTP 200",
     "POST /api/internship-batches/:id/import-participants | file Excel đúng định dạng",
     "HTTP 200, số lượng bản ghi được import thành công",
     "P"),
    (57,
     "Xuất danh sách tham gia đợt ra Excel",
     "1. Đăng nhập với quyền Admin.\n2. Vào trang chi tiết đợt thực tập.\n3. Nhấn nút Xuất Excel.\n4. Chờ file được tạo và tải về.\n5. Kiểm tra nội dung file Excel.",
     "File Excel chứa danh sách tham gia được tải về, HTTP 200",
     "GET /api/internship-batches/:id/export | token admin",
     "HTTP 200, file Excel hợp lệ được tải về",
     "P"),
    (58,
     "Admin xóa đợt thực tập chưa có dữ liệu",
     "1. Đăng nhập với quyền Admin.\n2. Vào danh sách đợt thực tập.\n3. Tìm đợt chưa có sinh viên tham gia.\n4. Nhấn nút Xóa bên cạnh đợt.\n5. Xác nhận xóa trong hộp thoại.",
     "Đợt thực tập bị xóa, HTTP 200",
     "DELETE /api/internship-batches/:id | token admin",
     "HTTP 200, đợt thực tập bị xóa khỏi CSDL",
     "P"),

    # ─── MODULE 10: BÁO CÁO THỰC TẬP (5 TC) ────────────────────────────────
    ("MODULE", "10", "Chức năng Báo cáo thực tập"),
    (59,
     "Sinh viên nộp báo cáo thực tập theo tuần",
     "1. Đăng nhập với tài khoản sinh viên.\n2. Vào trang Báo cáo thực tập.\n3. Nhấn nút Tạo báo cáo tuần.\n4. Nhập tiêu đề, nội dung công việc đã thực hiện.\n5. Nhấn Nộp báo cáo và kiểm tra trạng thái.",
     "Báo cáo được tạo, trạng thái: chưa duyệt, HTTP 201",
     "phan_cong_id | loai_bao_cao: tuan | tieu_de: Tuần 1 | noi_dung: Làm việc với...",
     "HTTP 201, báo cáo tuần được lưu, trạng thái: chưa duyệt",
     "P"),
    (60,
     "Sinh viên nộp báo cáo cuối kỳ thực tập",
     "1. Đăng nhập với tài khoản sinh viên.\n2. Vào trang Báo cáo thực tập.\n3. Nhấn nút Nộp báo cáo cuối kỳ.\n4. Nhập nội dung tổng kết và đính kèm file PDF.\n5. Nhấn Nộp và kiểm tra kết quả.",
     "Báo cáo cuối kỳ được tạo kèm file, HTTP 201",
     "phan_cong_id | loai_bao_cao: cuoi-khoa | file: BaoCaoTotNghiep.pdf",
     "HTTP 201, báo cáo cuối kỳ được lưu với file đính kèm",
     "P"),
    (61,
     "Giảng viên chấm điểm và nhận xét báo cáo",
     "1. Đăng nhập với tài khoản giảng viên.\n2. Vào trang Chấm báo cáo.\n3. Tìm báo cáo của SV cần chấm.\n4. Nhập điểm số: 8.5 và nhận xét.\n5. Nhấn Lưu và kiểm tra trạng thái báo cáo.",
     "Điểm và nhận xét được lưu, trạng thái: đã duyệt, HTTP 200",
     "POST /api/reports/weekly/:id/grade | diem: 8.5 | nhan_xet_gv: Làm tốt",
     "HTTP 200, điểm và nhận xét lưu, trạng thái báo cáo: đã duyệt",
     "P"),
    (62,
     "Giảng viên yêu cầu sinh viên sửa báo cáo",
     "1. Đăng nhập với tài khoản giảng viên.\n2. Vào trang Chấm báo cáo.\n3. Mở báo cáo chưa đạt yêu cầu.\n4. Nhập nhận xét: 'Cần bổ sung thêm...'\n5. Chọn trạng thái Cần sửa và nhấn Lưu.",
     "Trạng thái chuyển sang: cần sửa, SV nhận thông báo, HTTP 200",
     "PUT /api/reports/weekly/:id | trang_thai: can-sua | nhan_xet: Cần bổ sung thêm...",
     "HTTP 200, trạng thái: cần sửa, SV nhận thông báo yêu cầu sửa",
     "P"),
    (63,
     "Xem thống kê tổng quan báo cáo",
     "1. Đăng nhập (Admin hoặc Giảng viên).\n2. Vào trang Dashboard hoặc Báo cáo.\n3. Quan sát khung thống kê báo cáo.\n4. Kiểm tra số lượng theo từng trạng thái.\n5. So sánh với dữ liệu thực tế.",
     "Trả về số lượng báo cáo theo từng trạng thái, HTTP 200",
     "GET /api/reports/statistics | token admin",
     "HTTP 200, số lượng chưa duyệt, đã duyệt, cần sửa chính xác",
     "P"),

    # ─── MODULE 11: ĐỢT NỘP BÁO CÁO & BÀI NỘP (7 TC) ──────────────────────
    ("MODULE", "11", "Chức năng Đợt nộp báo cáo và bài nộp của sinh viên"),
    (64,
     "Giảng viên tạo đợt nộp báo cáo cho sinh viên",
     "1. Đăng nhập với tài khoản giảng viên.\n2. Vào trang Quản lý đợt nộp.\n3. Nhấn nút Tạo đợt nộp mới.\n4. Nhập tiêu đề, loại báo cáo, thời hạn nộp.\n5. Nhấn Lưu và kiểm tra đợt đã tạo.",
     "Đợt nộp được tạo, sinh viên có thể nộp bài, HTTP 201",
     "POST /api/teacher/submissions/slots | tieu_de: Tuần 3 | loai: tuan | end_at: 2024-12-20",
     "HTTP 201, đợt nộp báo cáo được tạo với thông tin đầy đủ",
     "P"),
    (65,
     "Sinh viên nộp bài trong thời hạn quy định",
     "1. Đăng nhập với tài khoản sinh viên.\n2. Vào trang Nộp báo cáo.\n3. Chọn đợt nộp đang mở.\n4. Nhấn Chọn file và chọn file PDF báo cáo.\n5. Nhấn Nộp bài và kiểm tra trạng thái.",
     "Bài nộp được lưu, trạng thái: đã nộp, HTTP 201",
     "POST /api/student/slots/:slotId/upload | file: BaoCaoTuan3.pdf (trước end_at)",
     "HTTP 201, bài nộp được lưu, trang_thai: đã nộp",
     "P"),
    (66,
     "Từ chối nộp bài khi hết thời hạn",
     "1. Đăng nhập với tài khoản sinh viên.\n2. Vào trang Nộp báo cáo.\n3. Chọn đợt nộp đã quá hạn (end_at đã qua).\n4. Chọn file và nhấn Nộp bài.\n5. Quan sát thông báo lỗi trả về.",
     "Hệ thống trả về lỗi, không cho nộp, HTTP 400",
     "POST /api/student/slots/:slotId/upload | sau end_at",
     "HTTP 400, message: 'Đã hết thời gian nộp bài'",
     "P"),
    (67,
     "Giảng viên nhận xét bài nộp của sinh viên",
     "1. Đăng nhập với tài khoản giảng viên.\n2. Vào trang Quản lý đợt nộp.\n3. Mở đợt nộp, chọn bài nộp của SV.\n4. Nhập nhận xét và chọn trạng thái đã chấm.\n5. Nhấn Lưu và kiểm tra kết quả.",
     "Nhận xét được lưu vào CSDL, HTTP 200",
     "POST /api/submissions/:submissionId/review | teacher_comment: Làm tốt | trang_thai: da_cham",
     "HTTP 200, nhận xét GV được lưu vào CSDL",
     "P"),
    (68,
     "Giảng viên xem trạng thái nộp bài trong đợt",
     "1. Đăng nhập với tài khoản giảng viên.\n2. Vào trang Quản lý đợt nộp.\n3. Chọn đợt nộp cần kiểm tra.\n4. Quan sát danh sách SV và trạng thái nộp.\n5. Xác nhận SV nào đã nộp, chưa nộp.",
     "Danh sách SV và trạng thái nộp được hiển thị, HTTP 200",
     "GET /api/teacher/submissions/slots/:slotId/statuses | token giang-vien",
     "HTTP 200, danh sách SV kèm trạng thái đã nộp/chưa nộp",
     "P"),
    (69,
     "Sinh viên xem lịch sử bài nộp của mình",
     "1. Đăng nhập với tài khoản sinh viên.\n2. Vào trang Lịch sử nộp bài.\n3. Quan sát danh sách các bài đã nộp.\n4. Kiểm tra thông tin: tên file, ngày nộp, trạng thái.\n5. Nhấn vào bài nộp để xem chi tiết.",
     "Lịch sử bài nộp đầy đủ được hiển thị, HTTP 200",
     "GET /api/student/all-my-submissions | token sinh-vien",
     "HTTP 200, danh sách các bài đã nộp với thông tin chi tiết",
     "P"),
    (70,
     "Xuất bảng điểm đợt nộp ra file Excel",
     "1. Đăng nhập với tài khoản giảng viên.\n2. Vào trang chi tiết đợt nộp.\n3. Nhấn nút Xuất bảng điểm.\n4. Chờ file được tạo và tải về.\n5. Mở file Excel kiểm tra nội dung.",
     "File Excel bảng điểm được tải về, HTTP 200",
     "GET /api/report-batches/:id/export | token giang-vien",
     "HTTP 200, file Excel bảng điểm hợp lệ được tải về",
     "P"),

    # ─── MODULE 12: ĐÁNH GIÁ THỰC TẬP (4 TC) ───────────────────────────────
    ("MODULE", "12", "Chức năng Đánh giá thực tập (GV và Doanh nghiệp)"),
    (71,
     "Doanh nghiệp đánh giá sinh viên thực tập",
     "1. Đăng nhập với tài khoản doanh nghiệp.\n2. Vào trang Danh sách SV thực tập.\n3. Chọn SV cần đánh giá.\n4. Nhập điểm số: 8.0 và nhận xét.\n5. Nhấn Lưu đánh giá.",
     "Đánh giá được lưu vào CSDL, HTTP 200",
     "PUT /api/company-internships/students/:studentId/evaluation | diem_so: 8.0 | nhan_xet: Chuẩn bị tốt",
     "HTTP 200, đánh giá của DN được lưu vào CSDL",
     "P"),
    (72,
     "Doanh nghiệp nộp toàn bộ đánh giá",
     "1. Đăng nhập với tài khoản doanh nghiệp.\n2. Hoàn thành đánh giá tất cả SV.\n3. Nhấn nút Nộp tất cả đánh giá.\n4. Xác nhận trong hộp thoại.\n5. Kiểm tra trạng thái phân công.",
     "Trạng thái phân công chuyển sang HOÀN_THÀNH, HTTP 200",
     "POST /api/company-internships/submit-all-evaluations | token doanh-nghiep",
     "HTTP 200, trạng thái phân công: HOÀN_THÀNH",
     "P"),
    (73,
     "Giảng viên nhập điểm cuối kỳ cho sinh viên",
     "1. Đăng nhập với tài khoản giảng viên.\n2. Vào trang Tổng hợp điểm.\n3. Tìm SV cần nhập điểm cuối kỳ.\n4. Nhập điểm giảng viên: 8.5.\n5. Nhấn Lưu và kiểm tra dữ liệu.",
     "Điểm GV được lưu vào CSDL, HTTP 200",
     "POST /api/teacher/reports/students/:maSinhVien/grade | diem_giang_vien: 8.5",
     "HTTP 200, điểm giảng viên được lưu vào bảng phan_cong_thuc_tap",
     "P"),
    (74,
     "Giảng viên xem bảng điểm tổng hợp",
     "1. Đăng nhập với tài khoản giảng viên.\n2. Vào trang Tổng hợp điểm.\n3. Quan sát bảng điểm của tất cả SV hướng dẫn.\n4. Kiểm tra cột điểm GV và điểm DN.\n5. Xác nhận dữ liệu chính xác.",
     "Bảng điểm tổng hợp với cả điểm GV và DN được hiển thị, HTTP 200",
     "GET /api/teacher/reports/grades-summary | token giang-vien",
     "HTTP 200, bảng điểm tổng hợp với điểm GV và điểm DN",
     "P"),

    # ─── MODULE 13: THÔNG BÁO HỆ THỐNG (4 TC) ──────────────────────────────
    ("MODULE", "13", "Chức năng Thông báo hệ thống"),
    (75,
     "Người dùng lấy danh sách thông báo cá nhân",
     "1. Đăng nhập vào hệ thống.\n2. Nhấn vào biểu tượng chuông thông báo.\n3. Quan sát danh sách thông báo.\n4. Kiểm tra phân loại chưa đọc/đã đọc.\n5. Xác nhận số lượng thông báo chưa đọc.",
     "Danh sách thông báo chưa đọc và đã đọc được hiển thị, HTTP 200",
     "GET /api/notifications/me | token hợp lệ",
     "HTTP 200, danh sách thông báo phân loại chưa đọc/đã đọc",
     "P"),
    (76,
     "Đánh dấu một thông báo là đã đọc",
     "1. Đăng nhập vào hệ thống.\n2. Mở danh sách thông báo.\n3. Nhấn vào một thông báo chưa đọc.\n4. Quan sát thay đổi trạng thái.\n5. Xác nhận thông báo chuyển sang đã đọc.",
     "Thông báo được đánh dấu là đã đọc, HTTP 200",
     "PATCH /api/notifications/:id | is_read: 1",
     "HTTP 200, thông báo được cập nhật is_read = 1",
     "P"),
    (77,
     "Đánh dấu tất cả thông báo là đã đọc",
     "1. Đăng nhập vào hệ thống.\n2. Mở danh sách thông báo.\n3. Nhấn nút Đọc tất cả.\n4. Quan sát danh sách thay đổi.\n5. Xác nhận không còn thông báo chưa đọc.",
     "Tất cả thông báo chưa đọc chuyển sang đã đọc, HTTP 200",
     "PATCH /api/notifications/bulk-update | action: mark_all_read",
     "HTTP 200, tất cả thông báo chưa đọc được cập nhật is_read = 1",
     "P"),
    (78,
     "Xóa một thông báo khỏi danh sách",
     "1. Đăng nhập vào hệ thống.\n2. Mở danh sách thông báo.\n3. Tìm thông báo cần xóa.\n4. Nhấn nút Xóa bên cạnh thông báo.\n5. Xác nhận xóa và kiểm tra danh sách.",
     "Thông báo bị xóa khỏi danh sách, HTTP 200",
     "DELETE /api/notifications/:id | token hợp lệ",
     "HTTP 200, thông báo bị xóa khỏi CSDL",
     "P"),

    # ─── MODULE 14: IMPORT DỮ LIỆU TỪ EXCEL (4 TC) ─────────────────────────
    ("MODULE", "14", "Chức năng Import dữ liệu từ Excel"),
    (79,
     "Admin tải file mẫu Excel để import sinh viên",
     "1. Đăng nhập với quyền Admin.\n2. Vào trang Import dữ liệu.\n3. Nhấn nút Tải file mẫu.\n4. Chờ file tải về máy tính.\n5. Mở file kiểm tra cấu trúc cột.",
     "File Excel mẫu được tải về máy tính, HTTP 200",
     "GET /api/import/template | token admin",
     "HTTP 200, file Excel mẫu có đúng cấu trúc được tải về",
     "P"),
    (80,
     "Kiểm tra dữ liệu file Excel trước khi import",
     "1. Đăng nhập với quyền Admin.\n2. Vào trang Import dữ liệu.\n3. Nhấn Chọn file và chọn file Excel.\n4. Nhấn nút Kiểm tra dữ liệu.\n5. Xem báo cáo lỗi và số bản ghi hợp lệ.",
     "Trả về danh sách lỗi (nếu có) và số bản ghi hợp lệ, HTTP 200",
     "POST /api/import/validate | file Excel danh sách SV",
     "HTTP 200, báo cáo validation: số bản ghi hợp lệ / lỗi",
     "P"),
    (81,
     "Import danh sách sinh viên từ file Excel hợp lệ",
     "1. Đăng nhập với quyền Admin.\n2. Vào trang Import sinh viên.\n3. Chọn file Excel đúng định dạng.\n4. Nhấn nút Import.\n5. Kiểm tra báo cáo số bản ghi thành công/thất bại.",
     "Dữ liệu được import, trả về số bản ghi thành công/thất bại, HTTP 200",
     "POST /api/import/sinh-vien | file Excel đúng định dạng",
     "HTTP 200, số bản ghi thành công và thất bại được báo cáo",
     "P"),
    (82,
     "Từ chối import file sai định dạng",
     "1. Đăng nhập với quyền Admin.\n2. Vào trang Import sinh viên.\n3. Chọn file PDF (không phải Excel).\n4. Nhấn nút Import.\n5. Quan sát thông báo lỗi trả về.",
     "Hệ thống trả về lỗi, không thực hiện import, HTTP 400",
     "POST /api/import/sinh-vien | file: DanhSach.pdf",
     "HTTP 400, message: 'Chỉ chấp nhận file Excel (.xlsx/.xls)'",
     "P"),

    # ─── MODULE 15: DASHBOARD & THỐNG KÊ (4 TC) ─────────────────────────────
    ("MODULE", "15", "Chức năng Dashboard và thống kê"),
    (83,
     "Admin xem dashboard tổng quan hệ thống",
     "1. Đăng nhập với quyền Admin.\n2. Vào trang Dashboard chính.\n3. Quan sát các khung thống kê.\n4. Kiểm tra số liệu: SV, GV, DN, đợt thực tập.\n5. Xác nhận dữ liệu khớp với CSDL.",
     "Hiển thị đầy đủ số liệu: SV, GV, DN, đợt thực tập, HTTP 200",
     "GET /api/dashboard/stats | token admin",
     "HTTP 200, trả về tổng SV, GV, DN, số đợt đang mở, số đơn chờ duyệt",
     "P"),
    (84,
     "Lấy thống kê cơ bản không cần đăng nhập",
     "1. Mở trình duyệt, không đăng nhập.\n2. Truy cập endpoint /api/dashboard/stats-public.\n3. Quan sát dữ liệu trả về.\n4. Kiểm tra không yêu cầu token.\n5. Xác nhận dữ liệu thống kê cơ bản.",
     "Trả về thông tin thống kê cơ bản, HTTP 200",
     "GET /api/dashboard/stats-public (không cần token)",
     "HTTP 200, thống kê cơ bản được trả về",
     "P"),
    (85,
     "Xem phân bổ sinh viên theo từng đợt thực tập",
     "1. Đăng nhập với quyền Admin.\n2. Vào trang Dashboard.\n3. Kéo đến phần Phân bổ sinh viên theo đợt.\n4. Quan sát biểu đồ/bảng phân bổ.\n5. Kiểm tra số liệu từng đợt.",
     "Trả về dữ liệu phân bổ SV theo đợt, HTTP 200",
     "GET /api/dashboard/students-by-period | token admin",
     "HTTP 200, dữ liệu phân bổ SV theo từng đợt thực tập",
     "P"),
    (86,
     "Xuất báo cáo sinh viên theo đợt ra Excel",
     "1. Đăng nhập với quyền Admin.\n2. Vào trang Dashboard.\n3. Nhấn nút Xuất Excel báo cáo SV theo đợt.\n4. Chờ file được tạo.\n5. Kiểm tra file Excel tải về.",
     "File Excel báo cáo SV theo đợt được tải về, HTTP 200",
     "GET /api/dashboard/export-students-by-period | token admin",
     "HTTP 200, file Excel báo cáo hợp lệ được tải về",
     "P"),

    # ─── MODULE 16: TÍCH HỢP ZALO BOT (2 TC) ───────────────────────────────
    ("MODULE", "16", "Chức năng Tích hợp Zalo Bot"),
    (87,
     "Gửi tin nhắn Zalo thông báo cho sinh viên",
     "1. Đăng nhập với quyền Admin.\n2. Vào trang Zalo Bot.\n3. Chọn tab Sinh viên.\n4. Chọn danh sách SV cần gửi, nhập nội dung tin nhắn.\n5. Nhấn Gửi và kiểm tra kết quả.",
     "Tin nhắn được gửi đến số điện thoại SV qua Flask/Zalo OA, HTTP 200",
     "POST /api/zalo/send-to-students | danh_sach_sv_id | noi_dung: 'Nhắc nộp báo cáo'",
     "HTTP 200, tin nhắn Zalo được gửi đến SV thành công",
     "P"),
    (88,
     "Gửi tin nhắn Zalo thông báo cho giảng viên",
     "1. Đăng nhập với quyền Admin.\n2. Vào trang Zalo Bot.\n3. Chọn tab Giảng viên.\n4. Chọn danh sách GV cần gửi, nhập nội dung tin nhắn.\n5. Nhấn Gửi và kiểm tra kết quả.",
     "Tin nhắn được gửi đến số điện thoại GV qua Flask/Zalo OA, HTTP 200",
     "POST /api/zalo/send-to-lecturers | danh_sach_gv_id | noi_dung: 'Yêu cầu chấm điểm'",
     "HTTP 200, tin nhắn Zalo được gửi đến GV thành công",
     "P"),

    # ─── MODULE 17: HỒ SƠ CÁ NHÂN (4 TC) ───────────────────────────────────
    ("MODULE", "17", "Chức năng Hồ sơ cá nhân"),
    (89,
     "Người dùng xem hồ sơ cá nhân",
     "1. Đăng nhập vào hệ thống (bất kỳ vai trò).\n2. Nhấn vào avatar hoặc tên người dùng.\n3. Chọn mục Hồ sơ cá nhân.\n4. Quan sát thông tin hiển thị.\n5. Kiểm tra các trường dữ liệu theo vai trò.",
     "Hiển thị đầy đủ thông tin cá nhân theo vai trò, HTTP 200",
     "GET /api/profile/me | token hợp lệ (bất kỳ vai trò nào)",
     "HTTP 200, trả về thông tin cá nhân đúng với vai trò đang đăng nhập",
     "P"),
    (90,
     "Người dùng cập nhật thông tin cá nhân",
     "1. Đăng nhập vào hệ thống.\n2. Vào trang Hồ sơ cá nhân.\n3. Nhấn nút Chỉnh sửa.\n4. Sửa số điện thoại: 0987654321 và địa chỉ: Hà Nội.\n5. Nhấn Lưu và kiểm tra thông tin cập nhật.",
     "Thông tin được cập nhật trong CSDL, HTTP 200",
     "PUT /api/profile/me | so_dien_thoai: 0987654321 | dia_chi: Hà Nội",
     "HTTP 200, thông tin cá nhân được cập nhật thành công",
     "P"),
    (91,
     "Người dùng upload ảnh đại diện mới",
     "1. Đăng nhập vào hệ thống.\n2. Vào trang Hồ sơ cá nhân.\n3. Nhấn vào biểu tượng camera trên ảnh đại diện.\n4. Chọn file ảnh JPG/PNG dưới 5MB.\n5. Xác nhận upload và kiểm tra ảnh mới trên Profile và Navbar.",
     "Ảnh đại diện được cập nhật hiển thị trên Profile và Navbar, HTTP 200",
     "POST /api/profile/avatar | Multipart file: avatar.jpg (dưới 5MB, JPG/PNG)",
     "HTTP 200, avatar_url được cập nhật, ảnh mới hiển thị trên Profile và Navbar",
     "P"),
    (92,
     "Người dùng xóa ảnh đại diện",
     "1. Đăng nhập vào hệ thống.\n2. Vào trang Hồ sơ cá nhân (đã có ảnh đại diện).\n3. Nhấn vào biểu tượng thùng rác trên ảnh.\n4. Xác nhận xóa trong hộp thoại.\n5. Kiểm tra avatar mặc định hiển thị.",
     "Ảnh đại diện bị xóa, hiển thị avatar mặc định (chữ cái đầu tên), HTTP 200",
     "DELETE /api/profile/avatar | token hợp lệ",
     "HTTP 200, avatar_url = NULL, hiển thị avatar mặc định",
     "P"),
]

# ── Helper functions ──────────────────────────────────────────────────────────
def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(shd)

def set_borders(cell, color="4472C4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def write_cell(cell, text, bold=False, center=False, size=9,
               color=None, italic=False):
    cell.text = ""
    align = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    lines = str(text).split("\n")
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.alignment = align
        r = p.add_run(line)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        r.font.name = "Times New Roman"
        if color:
            r.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def set_row_height(row, height_cm):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(height_cm * 567)))
    trHeight.set(qn("w:hRule"), "atLeast")
    trPr.append(trHeight)

# ── Main ──────────────────────────────────────────────────────────────────────
def build(output_path: str):
    doc = Document()

    # A4 ngang (landscape)
    sec = doc.sections[0]
    sec.page_width   = Cm(29.7)
    sec.page_height  = Cm(21.0)
    sec.left_margin  = Cm(1.5)
    sec.right_margin = Cm(1.5)
    sec.top_margin   = Cm(1.5)
    sec.bottom_margin = Cm(1.5)

    # ── Tiêu đề ─────────────────────────────────────────────────────────────
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("BẢNG TỔNG QUAN CÁC TEST CASE HỆ THỐNG")
    r.bold = True; r.font.size = Pt(14)
    r.font.name = "Times New Roman"
    r.font.color.rgb = C_HEADER

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = s.add_run("Hệ thống Quản lý Thực tập và Hợp tác Doanh nghiệp - Khoa CNTT - ĐH Đại Nam")
    sr.italic = True; sr.font.size = Pt(11)
    sr.font.name = "Times New Roman"

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.LEFT
    ir = info.add_run(
        "Dự án: Hệ thống Quản lý Thực tập CNTT ĐH Đại Nam          "
        "Người kiểm thử: Nhóm phát triển          "
        "Ngày test: ___/___/______"
    )
    ir.font.size = Pt(10)
    ir.font.name = "Times New Roman"

    doc.add_paragraph()

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run("Bảng 3.2: Bảng chi tiết các test case toàn hệ thống")
    cr.bold = True; cr.font.size = Pt(11)
    cr.font.name = "Times New Roman"
    cr.font.color.rgb = C_CAPTION

    # ── Số hàng ─────────────────────────────────────────────────────────────
    n_rows = 2 + len(TESTCASES)

    # ── Độ rộng cột (9 cột) ─────────────────────────────────────────────────
    COL_W = [Cm(0.9), Cm(4.0), Cm(4.5), Cm(4.2), Cm(4.2), Cm(4.0), Cm(1.4), Cm(1.2), Cm(2.5)]

    tbl = doc.add_table(rows=n_rows, cols=9)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"

    # ── Hàng header 1 ────────────────────────────────────────────────────────
    hdr1 = tbl.rows[0]
    set_row_height(hdr1, 0.9)
    labels1 = [
        "STT", "Mục tiêu kiểm thử", "Mô tả thao tác kiểm thử",
        "Kết quả chờ đợi", "Dữ liệu kiểm thử", "Kết quả thực tế",
        "Kiểm thử lần 1", "", ""
    ]
    for j, (lbl, w) in enumerate(zip(labels1, COL_W)):
        c = hdr1.cells[j]
        c.width = w
        set_cell_bg(c, C_HEADER)
        set_borders(c, "FFFFFF")
        write_cell(c, lbl, bold=True, center=True, size=10, color=C_WHITE)

    # Merge 3 cột cuối của hàng 1 → "Kiểm thử lần 1"
    hdr1.cells[6].merge(hdr1.cells[7])
    hdr1.cells[6].merge(hdr1.cells[8])
    hdr1.cells[6].width = COL_W[6]
    set_cell_bg(hdr1.cells[6], C_HEADER)
    set_borders(hdr1.cells[6], "FFFFFF")
    write_cell(hdr1.cells[6], "Kiểm thử lần 1", bold=True, center=True, size=10, color=C_WHITE)

    # Merge hàng 1 và hàng 2 cho 6 cột đầu (rowspan=2)
    for j in range(6):
        tbl.rows[0].cells[j].merge(tbl.rows[1].cells[j])

    # ── Hàng header 2 (sub-header) ───────────────────────────────────────────
    hdr2 = tbl.rows[1]
    set_row_height(hdr2, 0.8)
    sub = ["", "", "", "", "", "", "Kết quả\n(P/F/N)", "Bug #", "Diễn giải lỗi"]
    for j in range(6, 9):
        c = hdr2.cells[j]
        c.width = COL_W[j]
        set_cell_bg(c, C_HEADER)
        set_borders(c, "FFFFFF")
        write_cell(c, sub[j], bold=True, center=True, size=9, color=C_WHITE)

    # ── Hàng dữ liệu ─────────────────────────────────────────────────────────
    data_row = 2
    odd = True

    for tc in TESTCASES:
        row = tbl.rows[data_row]
        set_row_height(row, 0.8)

        if tc[0] == "MODULE":
            row.cells[0].merge(row.cells[8])
            cell = row.cells[0]
            set_cell_bg(cell, C_MODULE_BG)
            set_borders(cell, "4472C4")
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(f"MODULE {tc[1]}: {tc[2]}")
            r.bold = True
            r.font.size = Pt(10)
            r.font.name = "Times New Roman"
            r.font.color.rgb = C_MODULE_FG
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            odd = True
        else:
            stt, ten_tc, mo_ta, kq_cho_doi, du_lieu, kq_thuc_te, pfn = tc
            bg = C_ODD if odd else None
            vals    = [stt, ten_tc, mo_ta, kq_cho_doi, du_lieu, kq_thuc_te, pfn, "", ""]
            centers = [True, False, False, False, False, False, True, True, False]

            for j, (val, w, is_center) in enumerate(zip(vals, COL_W, centers)):
                c = row.cells[j]
                c.width = w
                if bg:
                    set_cell_bg(c, bg)
                set_borders(c, "4472C4")
                fc = C_PASS if (j == 6 and val == "P") else None
                write_cell(c, val,
                           center=is_center,
                           size=9,
                           color=fc,
                           bold=(j == 6 and val == "P"))
            odd = not odd

        data_row += 1

    doc.save(output_path)
    # Dùng ASCII để tránh lỗi encoding trên Windows
    print("[OK] Da xuat file Word thanh cong!")
    print(f"[PATH] {output_path}")

if __name__ == "__main__":
    out = os.path.join(
        os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
        "backend", "docs", "Bang_TestCase_ChiTiet_v2.docx"
    )
    build(out)
