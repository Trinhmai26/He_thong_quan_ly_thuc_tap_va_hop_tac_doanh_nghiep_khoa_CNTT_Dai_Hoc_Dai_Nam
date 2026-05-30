# -*- coding: utf-8 -*-
"""
Xuat bang tong quan testcase (kieu anh mau) ra file Word.
Format: STT | Muc tieu kiem thu | Tong so (Pass / Fail)
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Dữ liệu tổng quan từng chức năng ─────────────────────────────────────────
# (stt, ten_chuc_nang, so_pass, so_fail)
SUMMARY = [
    (1,  "Chức năng Xác thực hệ thống (Đăng nhập / Đăng xuất)",         9,  0),
    (2,  "Chức năng Quản lý tài khoản và mật khẩu",                      9,  0),
    (3,  "Chức năng Quản lý sinh viên",                                   9,  0),
    (4,  "Chức năng Quản lý giảng viên",                                  5,  0),
    (5,  "Chức năng Quản lý doanh nghiệp",                                5,  0),
    (6,  "Chức năng Đăng ký thực tập",                                    6,  0),
    (7,  "Chức năng Quy trình phỏng vấn (Interview Workflow)",            5,  0),
    (8,  "Chức năng Phân công thực tập",                                  4,  0),
    (9,  "Chức năng Quản lý đợt thực tập",                               6,  0),
    (10, "Chức năng Báo cáo thực tập",                                    5,  0),
    (11, "Chức năng Đợt nộp báo cáo và bài nộp của sinh viên",           7,  0),
    (12, "Chức năng Đánh giá thực tập (GV và Doanh nghiệp)",             4,  0),
    (13, "Chức năng Thông báo hệ thống",                                  4,  0),
    (14, "Chức năng Import dữ liệu từ Excel",                             4,  0),
    (15, "Chức năng Dashboard và thống kê",                               4,  0),
    (16, "Chức năng Tích hợp Zalo Bot",                                   2,  0),
    (17, "Chức năng Hồ sơ cá nhân",                                       4,  0),
]

TOTAL_PASS = sum(r[2] for r in SUMMARY)
TOTAL_FAIL = sum(r[3] for r in SUMMARY)

# ── Màu sắc ──────────────────────────────────────────────────────────────────
C_HEADER    = RGBColor(0x1F, 0x49, 0x7D)
C_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
C_ODD       = RGBColor(0xED, 0xF3, 0xFB)
C_TOTAL_BG  = RGBColor(0xD6, 0xE4, 0xF7)
C_PASS      = RGBColor(0x00, 0x70, 0x00)
C_FAIL      = RGBColor(0xC0, 0x00, 0x00)
C_CAPTION   = RGBColor(0x1F, 0x49, 0x7D)
C_BLACK     = RGBColor(0x00, 0x00, 0x00)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(shd)

def set_borders(cell, color="808080"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement("w:tcBorders")
    for side in ("top","left","bottom","right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tcB.append(b)
    tcPr.append(tcB)

def write(cell, text, bold=False, center=False, size=11,
          color: RGBColor | None = None, italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(str(text))
    r.bold = bold; r.italic = italic
    r.font.size = Pt(size)
    r.font.name = "Times New Roman"
    r.font.color.rgb = color if color else C_BLACK
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def set_col_width(table, col_idx, width):
    for row in table.rows:
        row.cells[col_idx].width = width

# ── Build ─────────────────────────────────────────────────────────────────────
def build(output_path: str):
    doc = Document()

    # Trang A4 đứng
    sec = doc.sections[0]
    sec.page_width  = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin   = Cm(3)
    sec.right_margin  = Cm(2)
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)

    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run("Bảng 3.1: Bảng tổng quan về các testcase thực hiện")
    cr.bold = True; cr.font.size = Pt(12)
    cr.font.name = "Times New Roman"
    cr.font.color.rgb = C_CAPTION

    doc.add_paragraph()

    # Tổng số hàng = 2 header + n data + 1 tổng cộng
    n = len(SUMMARY)
    tbl = doc.add_table(rows=2 + n + 1, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"

    # ── Hàng 0: header chính (STT | Mục tiêu | Tổng số <merged 2 col>) ──────
    r0 = tbl.rows[0].cells

    # Merge hàng 0 cột 0 với hàng 1 cột 0 (STT)
    r0[0].merge(tbl.rows[1].cells[0])
    set_bg(r0[0], C_HEADER); set_borders(r0[0], "FFFFFF")
    write(r0[0], "STT", bold=True, center=True, size=11, color=C_WHITE)

    # Merge hàng 0 cột 1 với hàng 1 cột 1 (Mục tiêu kiểm thử)
    r0[1].merge(tbl.rows[1].cells[1])
    set_bg(r0[1], C_HEADER); set_borders(r0[1], "FFFFFF")
    write(r0[1], "Mục tiêu kiểm thử", bold=True, center=True, size=11, color=C_WHITE)

    # Merge hàng 0 cột 2 + cột 3 → "Tổng số"
    r0[2].merge(r0[3])
    set_bg(r0[2], C_HEADER); set_borders(r0[2], "FFFFFF")
    write(r0[2], "Tổng số", bold=True, center=True, size=11, color=C_WHITE)

    # ── Hàng 1: sub-header (Pass | Fail) ─────────────────────────────────────
    r1 = tbl.rows[1].cells
    # cells[0] và [1] đã bị merge ở trên → chỉ xử lý [2] và [3]
    for ci, label in [(2, "Pass"), (3, "Fail")]:
        set_bg(r1[ci], C_HEADER); set_borders(r1[ci], "FFFFFF")
        write(r1[ci], label, bold=True, center=True, size=11, color=C_WHITE)

    # ── Hàng data ─────────────────────────────────────────────────────────────
    for i, (stt, ten, n_pass, n_fail) in enumerate(SUMMARY):
        row = tbl.rows[2 + i].cells
        bg = C_ODD if i % 2 == 0 else None

        # STT
        if bg: set_bg(row[0], bg)
        set_borders(row[0], "B0B0B0")
        write(row[0], stt, center=True, size=11)

        # Tên chức năng
        if bg: set_bg(row[1], bg)
        set_borders(row[1], "B0B0B0")
        write(row[1], ten, size=11)

        # Pass
        if bg: set_bg(row[2], bg)
        set_borders(row[2], "B0B0B0")
        write(row[2], n_pass, center=True, size=11,
              color=C_PASS if n_pass > 0 else C_BLACK, bold=(n_pass > 0))

        # Fail
        if bg: set_bg(row[3], bg)
        set_borders(row[3], "B0B0B0")
        write(row[3], n_fail, center=True, size=11,
              color=C_FAIL if n_fail > 0 else C_BLACK, bold=(n_fail > 0))

    # ── Hàng tổng cộng ────────────────────────────────────────────────────────
    tot = tbl.rows[2 + n].cells
    for c in tot:
        set_bg(c, C_TOTAL_BG)
        set_borders(c, "808080")

    write(tot[0], "Tổng", bold=True, center=True, size=11, color=C_HEADER)
    # Merge cột 0 + 1 cho "Tổng"
    tot[0].merge(tot[1])
    set_bg(tot[0], C_TOTAL_BG); set_borders(tot[0], "808080")
    write(tot[0], "Tổng", bold=True, center=True, size=11, color=C_HEADER)

    write(tot[2], TOTAL_PASS, bold=True, center=True, size=11, color=C_PASS)
    write(tot[3], TOTAL_FAIL, bold=True, center=True, size=11,
          color=C_FAIL if TOTAL_FAIL > 0 else C_BLACK)

    # ── Chiều rộng cột ────────────────────────────────────────────────────────
    col_widths = [Cm(1.5), Cm(10.0), Cm(2.5), Cm(2.0)]
    for row in tbl.rows:
        cells = row.cells
        seen = set()
        for j, w in enumerate(col_widths):
            tc_id = id(cells[j]._tc)
            if tc_id not in seen:
                cells[j].width = w
                seen.add(tc_id)

    doc.save(output_path)
    print(f"[OK] Da xuat: {output_path}")

if __name__ == "__main__":
    out = os.path.join(
        os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
        "backend", "docs", "Bang_Tong_Quan_TestCase_Summary.docx"
    )
    build(out)
