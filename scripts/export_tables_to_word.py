"""
Xuất 20 bảng MySQL của hệ thống quản lý thực tập ra file Word.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── Dữ liệu 20 bảng ──────────────────────────────────────────────────────────

TABLES = [
    {
        "index": "2.1",
        "name": "accounts",
        "title": "Bảng accounts (Tài khoản đăng nhập)",
        "columns": [
            (1, "id", "INT", "PK, NOT NULL, AUTO_INCREMENT", "Mã định danh tài khoản"),
            (2, "user_id", "VARCHAR(50)", "UNIQUE, NOT NULL", "Mã đăng nhập"),
            (3, "email", "VARCHAR(255)", "UNIQUE, NOT NULL", "Địa chỉ email đăng nhập"),
            (4, "password_hash", "VARCHAR(255)", "NOT NULL", "Mật khẩu đã mã hóa"),
            (5, "role", "ENUM", "NOT NULL", "Vai trò: admin, sinh-vien, giang-vien, doanh-nghiep"),
            (6, "is_active", "TINYINT(1)", "DEFAULT 1", "Trạng thái hoạt động của tài khoản"),
            (7, "last_login", "TIMESTAMP", "NULL", "Thời điểm đăng nhập gần nhất"),
            (8, "created_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP", "Thời điểm tạo tài khoản"),
            (9, "updated_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP ON UPDATE", "Thời điểm cập nhật gần nhất"),
        ],
    },
    {
        "index": "2.2",
        "name": "admin",
        "title": "Bảng admin (Quản trị viên)",
        "columns": [
            (1, "id", "INT", "PK, NOT NULL, AUTO_INCREMENT", "Mã định danh quản trị viên"),
            (2, "account_id", "INT", "NOT NULL, FK → accounts(id)", "Mã tài khoản liên kết"),
            (3, "full_name", "VARCHAR(255)", "NOT NULL", "Họ tên quản trị viên"),
            (4, "phone", "VARCHAR(20)", "NULL", "Số điện thoại liên hệ"),
            (5, "dia_chi", "VARCHAR(255)", "NULL", "Địa chỉ liên hệ"),
            (6, "position", "VARCHAR(100)", "NULL", "Chức vụ quản trị"),
            (7, "permissions", "JSON", "NULL", "Quyền hạn cụ thể"),
            (8, "created_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP", "Thời điểm tạo bản ghi"),
            (9, "updated_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP ON UPDATE", "Thời điểm cập nhật gần nhất"),
        ],
    },
    {
        "index": "2.3",
        "name": "sinh_vien",
        "title": "Bảng sinh_vien (Sinh viên)",
        "columns": [
            (1, "id", "INT", "PK, NOT NULL, AUTO_INCREMENT", "Mã định danh sinh viên"),
            (2, "account_id", "INT", "NOT NULL, FK → accounts(id)", "Mã tài khoản liên kết"),
            (3, "ma_sinh_vien", "VARCHAR(20)", "UNIQUE, NOT NULL", "Mã sinh viên"),
            (4, "ho_ten", "VARCHAR(255)", "NOT NULL", "Họ tên sinh viên"),
            (5, "lop", "VARCHAR(50)", "NULL", "Lớp học"),
            (6, "khoa", "VARCHAR(100)", "NULL", "Khoa"),
            (7, "nganh", "VARCHAR(100)", "NULL", "Ngành học"),
            (8, "khoa_hoc", "VARCHAR(20)", "NULL", "Khóa học (VD: K17, K18)"),
            (9, "ngay_sinh", "DATE", "NULL", "Ngày sinh"),
            (10, "gioi_tinh", "ENUM", "NULL", "Giới tính: Nam, Nữ, Khác"),
            (11, "dia_chi", "TEXT", "NULL", "Địa chỉ"),
            (12, "so_dien_thoai", "VARCHAR(20)", "NULL", "Số điện thoại"),
            (13, "email_ca_nhan", "VARCHAR(255)", "NULL", "Email cá nhân"),
            (14, "gpa", "DECIMAL(3,2)", "NULL", "Điểm trung bình tích lũy"),
            (15, "so_tc_tich_luy", "INT", "NULL", "Số tín chỉ tích lũy"),
            (16, "so_tc_ht", "INT", "NULL", "Số tín chỉ hoàn thành"),
            (17, "nam_thu", "INT", "NULL", "Năm thứ đang học"),
            (18, "hp_no", "INT", "NULL", "Học phần nợ"),
            (19, "tinh_trang_hoc_tap", "ENUM", "DEFAULT 'Đang học'", "Đang học, Tạm nghỉ, Thôi học, Tốt nghiệp"),
            (20, "giang_vien_huong_dan", "VARCHAR(255)", "NULL", "Giảng viên hướng dẫn"),
            (21, "nguyen_vong_thuc_tap", "VARCHAR(50)", "NULL", "Nguyện vọng thực tập"),
            (22, "vi_tri_muon_ung_tuyen_thuc_tap", "VARCHAR(255)", "NULL", "Vị trí muốn ứng tuyển thực tập"),
            (23, "don_vi_thuc_tap", "VARCHAR(255)", "NULL", "Đơn vị thực tập"),
            (24, "cong_ty_tu_lien_he", "VARCHAR(255)", "NULL", "Công ty tự liên hệ"),
            (25, "dia_chi_cong_ty", "TEXT", "NULL", "Địa chỉ công ty"),
            (26, "nguoi_lien_he_cong_ty", "VARCHAR(255)", "NULL", "Người liên hệ công ty"),
            (27, "sdt_nguoi_lien_he", "VARCHAR(20)", "NULL", "SĐT người liên hệ"),
            (28, "cv_path", "VARCHAR(500)", "NULL", "Đường dẫn file CV"),
            (29, "zalo_user_id", "VARCHAR(100)", "NULL", "ID Zalo của sinh viên"),
            (30, "trang_thai_phan_cong", "ENUM", "DEFAULT 'chua-phan-cong'", "Trạng thái phân công"),
            (31, "dot_thuc_tap_admin", "ENUM", "NULL", "Đợt thực tập: dot-1, dot-2"),
            (32, "dot_thuc_tap_id", "INT", "FK → dot_thuc_tap(id)", "Mã đợt thực tập"),
            (33, "created_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP", "Thời điểm tạo bản ghi"),
            (34, "updated_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP ON UPDATE", "Thời điểm cập nhật gần nhất"),
        ],
    },
    {
        "index": "2.4",
        "name": "giang_vien",
        "title": "Bảng giang_vien (Giảng viên)",
        "columns": [
            (1, "id", "INT", "PK, NOT NULL, AUTO_INCREMENT", "Mã định danh giảng viên"),
            (2, "account_id", "INT", "NOT NULL, FK → accounts(id)", "Mã tài khoản liên kết"),
            (3, "ma_giang_vien", "VARCHAR(20)", "UNIQUE, NOT NULL", "Mã giảng viên"),
            (4, "ho_ten", "VARCHAR(255)", "NOT NULL", "Họ tên giảng viên"),
            (5, "khoa", "VARCHAR(100)", "NOT NULL", "Khoa"),
            (6, "bo_mon", "VARCHAR(100)", "NULL", "Bộ môn"),
            (7, "chuc_vu", "VARCHAR(100)", "NULL", "Chức vụ"),
            (8, "hoc_vi", "VARCHAR(50)", "NULL", "Học vị (Thạc sĩ, Tiến sĩ, ...)"),
            (9, "chuyen_mon", "TEXT", "NULL", "Chuyên môn, lĩnh vực nghiên cứu"),
            (10, "so_dien_thoai", "VARCHAR(20)", "NULL", "Số điện thoại"),
            (11, "email_ca_nhan", "VARCHAR(255)", "NULL", "Email cá nhân"),
            (12, "dia_chi", "TEXT", "NULL", "Địa chỉ"),
            (13, "kinh_nghiem_lam_viec", "TEXT", "NULL", "Kinh nghiệm làm việc"),
            (14, "bang_cap", "TEXT", "NULL", "Các bằng cấp đã có"),
            (15, "ngay_sinh", "DATE", "NULL", "Ngày sinh"),
            (16, "chuc_danh", "VARCHAR(100)", "NULL", "Chức danh"),
            (17, "can_cuoc_cong_dan", "VARCHAR(20)", "NULL", "Số CCCD"),
            (18, "zalo_user_id", "VARCHAR(100)", "NULL", "ID Zalo của giảng viên"),
            (19, "created_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP", "Thời điểm tạo bản ghi"),
            (20, "updated_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP ON UPDATE", "Thời điểm cập nhật gần nhất"),
        ],
    },
    {
        "index": "2.5",
        "name": "doanh_nghiep",
        "title": "Bảng doanh_nghiep (Doanh nghiệp)",
        "columns": [
            (1, "id", "INT", "PK, NOT NULL, AUTO_INCREMENT", "Mã định danh doanh nghiệp"),
            (2, "account_id", "INT", "NOT NULL, FK → accounts(id)", "Mã tài khoản liên kết"),
            (3, "ma_doanh_nghiep", "VARCHAR(20)", "UNIQUE, NOT NULL", "Mã doanh nghiệp"),
            (4, "ten_cong_ty", "VARCHAR(255)", "NOT NULL", "Tên công ty"),
            (5, "ten_nguoi_lien_he", "VARCHAR(255)", "NOT NULL", "Tên người đại diện liên hệ"),
            (6, "chuc_vu_nguoi_lien_he", "VARCHAR(100)", "NULL", "Chức vụ người liên hệ"),
            (7, "dia_chi_cong_ty", "TEXT", "NOT NULL", "Địa chỉ công ty"),
            (8, "so_dien_thoai", "VARCHAR(20)", "NOT NULL", "Số điện thoại"),
            (9, "email_cong_ty", "VARCHAR(255)", "NULL", "Email công ty"),
            (10, "website", "VARCHAR(255)", "NULL", "Website công ty"),
            (11, "linh_vuc_hoat_dong", "VARCHAR(255)", "NULL", "Lĩnh vực kinh doanh"),
            (12, "quy_mo_nhan_su", "VARCHAR(50)", "NULL", "Quy mô nhân sự"),
            (13, "mo_ta_cong_ty", "TEXT", "NULL", "Mô tả công ty"),
            (14, "yeu_cau_thuc_tap", "TEXT", "NULL", "Yêu cầu đối với sinh viên thực tập"),
            (15, "so_luong_nhan_thuc_tap", "INT", "DEFAULT 0", "Số lượng sinh viên có thể nhận thực tập"),
            (16, "thoi_gian_thuc_tap", "VARCHAR(100)", "NULL", "Thời gian thực tập"),
            (17, "dia_chi_thuc_tap", "TEXT", "NULL", "Địa chỉ nơi thực tập"),
            (18, "trang_thai_hop_tac", "ENUM", "DEFAULT 'Đang hợp tác'", "Đang hợp tác, Tạm dừng, Ngừng hợp tác"),
            (19, "vi_tri_tuyen_dung", "TEXT", "NULL", "Vị trí tuyển dụng"),
            (20, "created_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP", "Thời điểm tạo bản ghi"),
            (21, "updated_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP ON UPDATE", "Thời điểm cập nhật gần nhất"),
        ],
    },
    {
        "index": "2.6",
        "name": "dot_thuc_tap",
        "title": "Bảng dot_thuc_tap (Đợt thực tập)",
        "columns": [
            (1, "id", "INT", "PK, NOT NULL, AUTO_INCREMENT", "Mã định danh đợt thực tập"),
            (2, "ten_dot", "VARCHAR(255)", "NOT NULL", "Tên đợt thực tập"),
            (3, "thoi_gian_bat_dau", "DATE", "NOT NULL", "Ngày bắt đầu đợt"),
            (4, "thoi_gian_ket_thuc", "DATE", "NOT NULL", "Ngày kết thúc đợt"),
            (5, "mo_ta", "TEXT", "NULL", "Mô tả đợt thực tập"),
            (6, "trang_thai", "ENUM", "DEFAULT 'sap-mo'", "Trạng thái: sap-mo, dang-dien-ra, ket-thuc"),
            (7, "so_sinh_vien_tham_gia", "INT", "NOT NULL, DEFAULT 0", "Số sinh viên tham gia"),
            (8, "so_giang_vien_huong_dan", "INT", "NOT NULL, DEFAULT 0", "Số giảng viên hướng dẫn"),
            (9, "so_doanh_nghiep_tham_gia", "INT", "NOT NULL, DEFAULT 0", "Số doanh nghiệp tham gia"),
            (10, "thoi_gian_dang_ky_tu", "DATE", "NULL", "Thời gian bắt đầu đăng ký"),
            (11, "thoi_gian_dang_ky_den", "DATE", "NULL", "Thời gian kết thúc đăng ký"),
            (12, "khoa_hoc_ap_dung", "VARCHAR(50)", "NULL", "Khóa học áp dụng"),
            (13, "lop_ap_dung", "VARCHAR(50)", "NULL", "Lớp áp dụng"),
            (14, "thoi_gian_thuc_tap_dot_1_tu", "DATE", "NULL", "Thời gian bắt đầu đợt nhỏ 1"),
            (15, "thoi_gian_thuc_tap_dot_1_den", "DATE", "NULL", "Thời gian kết thúc đợt nhỏ 1"),
            (16, "thoi_gian_thuc_tap_dot_2_tu", "DATE", "NULL", "Thời gian bắt đầu đợt nhỏ 2"),
            (17, "thoi_gian_thuc_tap_dot_2_den", "DATE", "NULL", "Thời gian kết thúc đợt nhỏ 2"),
            (18, "dot_nho_config", "LONGTEXT", "NULL", "Cấu hình đợt nhỏ (JSON)"),
            (19, "so_sinh_vien_excel", "INT", "DEFAULT 0", "Số sinh viên import từ Excel"),
            (20, "so_giang_vien_excel", "INT", "DEFAULT 0", "Số giảng viên import từ Excel"),
            (21, "so_doanh_nghiep_excel", "INT", "DEFAULT 0", "Số doanh nghiệp import từ Excel"),
            (22, "created_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP", "Thời điểm tạo bản ghi"),
            (23, "updated_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP ON UPDATE", "Thời điểm cập nhật gần nhất"),
        ],
    },
    {
        "index": "2.7",
        "name": "dang_ky_thuc_tap_sinh_vien",
        "title": "Bảng dang_ky_thuc_tap_sinh_vien (Đăng ký thực tập)",
        "columns": [
            (1, "id", "INT", "PK, NOT NULL, AUTO_INCREMENT", "Mã đăng ký"),
            (2, "sinh_vien_id", "INT", "NOT NULL, FK → sinh_vien(id)", "Mã sinh viên"),
            (3, "nguyen_vong_thuc_tap", "ENUM", "NOT NULL", "Nguyện vọng: khoa-gioi-thieu, tu-lien-he"),
            (4, "vi_tri_thuc_tap_mong_muon", "VARCHAR(255)", "NOT NULL", "Vị trí thực tập mong muốn"),
            (5, "ten_cong_ty", "VARCHAR(255)", "NULL", "Tên công ty tự liên hệ"),
            (6, "dia_chi_cong_ty", "TEXT", "NULL", "Địa chỉ công ty"),
            (7, "nguoi_lien_he", "VARCHAR(255)", "NULL", "Người liên hệ tại công ty"),
            (8, "so_dien_thoai_lien_he", "VARCHAR(20)", "NULL", "SĐT người liên hệ"),
            (9, "ghi_chu", "TEXT", "NULL", "Ghi chú bổ sung"),
            (10, "trang_thai", "ENUM", "DEFAULT 'cho-duyet'", "Trạng thái: cho-duyet, da-duyet, tu-choi"),
            (11, "workflow_status", "ENUM", "DEFAULT 'CHO_DUYET'", "Trạng thái quy trình chi tiết"),
            (12, "workflow_status_v2", "ENUM", "NOT NULL, DEFAULT 'PENDING'", "PENDING, APPROVED, REJECTED, INTERVIEW_SCHEDULED, PASS, FAIL"),
            (13, "ly_do_tu_choi", "TEXT", "NULL", "Lý do từ chối"),
            (14, "nguoi_duyet_id", "INT", "NULL", "Mã người duyệt"),
            (15, "ngay_duyet", "DATETIME", "NULL", "Ngày duyệt"),
            (16, "interview_date", "DATE", "NULL", "Ngày phỏng vấn"),
            (17, "interview_time", "TIME", "NULL", "Giờ phỏng vấn"),
            (18, "interview_location", "VARCHAR(255)", "NULL", "Địa điểm phỏng vấn"),
            (19, "interview_note", "TEXT", "NULL", "Ghi chú phỏng vấn"),
            (20, "interview_updated_at", "DATETIME", "NULL", "Thời điểm cập nhật lịch phỏng vấn"),
            (21, "result_note", "TEXT", "NULL", "Ghi chú kết quả"),
            (22, "created_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP", "Thời điểm tạo bản ghi"),
            (23, "updated_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP ON UPDATE", "Thời điểm cập nhật gần nhất"),
        ],
    },
    {
        "index": "2.8",
        "name": "phan_cong_thuc_tap",
        "title": "Bảng phan_cong_thuc_tap (Phân công thực tập)",
        "columns": [
            (1, "id", "INT", "PK, NOT NULL, AUTO_INCREMENT", "Mã phân công"),
            (2, "sinh_vien_id", "INT", "NOT NULL, FK → sinh_vien(id)", "Mã sinh viên được phân công"),
            (3, "doanh_nghiep_id", "INT", "NOT NULL, FK → doanh_nghiep(id)", "Mã doanh nghiệp thực tập"),
            (4, "dot_thuc_tap_id", "INT", "NOT NULL, FK → dot_thuc_tap(id)", "Mã đợt thực tập"),
            (5, "giang_vien_id", "INT", "FK → giang_vien(id)", "Mã giảng viên hướng dẫn"),
            (6, "ngay_bat_dau", "DATE", "NOT NULL", "Ngày bắt đầu thực tập"),
            (7, "ngay_ket_thuc", "DATE", "NOT NULL", "Ngày kết thúc thực tập"),
            (8, "trang_thai", "ENUM", "DEFAULT 'chua-bat-dau'", "chua-bat-dau, dang-dien-ra, hoan-thanh, tam-dung"),
            (9, "workflow_status", "ENUM", "DEFAULT 'DA_PHAN_CONG'", "Trạng thái quy trình"),
            (10, "workflow_updated_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP", "Thời điểm cập nhật workflow"),
            (11, "diem_so", "DECIMAL(3,1)", "NULL", "Điểm tổng kết"),
            (12, "nhan_xet", "TEXT", "NULL", "Nhận xét của doanh nghiệp"),
            (13, "ngay_nop_danh_gia", "TIMESTAMP", "NULL", "Ngày nộp đánh giá"),
            (14, "diem_giang_vien", "DECIMAL(4,2)", "NULL", "Điểm giảng viên chấm"),
            (15, "nhan_xet_giang_vien", "TEXT", "NULL", "Nhận xét của giảng viên"),
            (16, "created_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP", "Thời điểm tạo bản ghi"),
            (17, "updated_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP ON UPDATE", "Thời điểm cập nhật gần nhất"),
        ],
    },
    {
        "index": "2.9",
        "name": "tin_tuyen_dung",
        "title": "Bảng tin_tuyen_dung (Tin tuyển dụng)",
        "columns": [
            (1, "id", "INT", "PK, NOT NULL, AUTO_INCREMENT", "Mã tin tuyển dụng"),
            (2, "doanh_nghiep_id", "INT", "NOT NULL, FK → doanh_nghiep(id)", "Mã doanh nghiệp đăng tin"),
            (3, "tieu_de", "VARCHAR(255)", "NOT NULL", "Tiêu đề tin tuyển dụng"),
            (4, "mo_ta_cong_viec", "TEXT", "NOT NULL", "Mô tả công việc"),
            (5, "yeu_cau", "TEXT", "NOT NULL", "Yêu cầu ứng viên"),
            (6, "so_luong_tuyen", "INT", "DEFAULT 1", "Số lượng tuyển dụng"),
            (7, "muc_luong", "VARCHAR(100)", "NULL", "Mức lương/thù lao"),
            (8, "hinh_thuc_lam_viec", "VARCHAR(100)", "NULL", "Hình thức làm việc"),
            (9, "dia_diem", "VARCHAR(255)", "NULL", "Địa điểm làm việc"),
            (10, "han_ung_tuyen", "DATE", "NULL", "Hạn ứng tuyển"),
            (11, "trang_thai", "ENUM", "DEFAULT 'dang-tuyen'", "Trạng thái: dang-tuyen, tam-dung, het-han"),
            (12, "created_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP", "Thời điểm tạo bản ghi"),
            (13, "updated_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP ON UPDATE", "Thời điểm cập nhật gần nhất"),
        ],
    },
    {
        "index": "2.10",
        "name": "ung_tuyen",
        "title": "Bảng ung_tuyen (Ứng tuyển)",
        "columns": [
            (1, "id", "INT", "PK, NOT NULL, AUTO_INCREMENT", "Mã ứng tuyển"),
            (2, "sinh_vien_id", "INT", "UNIQUE, NOT NULL, FK → sinh_vien(id)", "Mã sinh viên ứng tuyển"),
            (3, "tin_tuyen_dung_id", "INT", "UNIQUE, NOT NULL, FK → tin_tuyen_dung(id)", "Mã tin tuyển dụng"),
            (4, "thu_xin_viec", "TEXT", "NOT NULL", "Thư xin việc"),
            (5, "cv_file", "VARCHAR(500)", "NULL", "Đường dẫn file CV"),
            (6, "ngay_ung_tuyen", "DATE", "NOT NULL", "Ngày ứng tuyển"),
            (7, "trang_thai", "ENUM", "DEFAULT 'dang-cho'", "Trạng thái: dang-cho, duoc-chap-nhan, bi-tu-choi"),
            (8, "ghi_chu", "TEXT", "NULL", "Ghi chú"),
            (9, "created_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP", "Thời điểm tạo bản ghi"),
            (10, "updated_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP ON UPDATE", "Thời điểm cập nhật gần nhất"),
        ],
    },
    {
        "index": "2.11",
        "name": "bao_cao_thuc_tap",
        "title": "Bảng bao_cao_thuc_tap (Báo cáo thực tập)",
        "columns": [
            (1, "id", "INT", "PK, NOT NULL, AUTO_INCREMENT", "Mã báo cáo"),
            (2, "phan_cong_id", "INT", "NOT NULL, FK → phan_cong_thuc_tap(id)", "Mã phân công thực tập"),
            (3, "loai_bao_cao", "ENUM", "NOT NULL", "Loại: tuan, thang, cuoi-khoa"),
            (4, "tieu_de", "VARCHAR(255)", "NOT NULL", "Tiêu đề báo cáo"),
            (5, "noi_dung", "TEXT", "NOT NULL", "Nội dung báo cáo"),
            (6, "file_dinh_kem", "VARCHAR(500)", "NULL", "Đường dẫn file đính kèm"),
            (7, "ngay_nop", "DATE", "NOT NULL", "Ngày nộp báo cáo"),
            (8, "trang_thai", "ENUM", "DEFAULT 'chua-duyet'", "Trạng thái: chua-duyet, da-duyet, can-sua"),
            (9, "nhan_xet_gv", "TEXT", "NULL", "Nhận xét của giảng viên"),
            (10, "created_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP", "Thời điểm tạo bản ghi"),
            (11, "updated_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP ON UPDATE", "Thời điểm cập nhật gần nhất"),
        ],
    },
    {
        "index": "2.12",
        "name": "dot_nop_bao_cao_theo_tuan",
        "title": "Bảng dot_nop_bao_cao_theo_tuan (Đợt nộp báo cáo)",
        "columns": [
            (1, "id", "INT", "PK, NOT NULL, AUTO_INCREMENT", "Mã đợt nộp"),
            (2, "ma_giang_vien", "VARCHAR(20)", "NOT NULL", "Mã giảng viên tạo đợt"),
            (3, "tieu_de", "VARCHAR(255)", "NOT NULL", "Tiêu đề đợt nộp"),
            (4, "loai_bao_cao", "ENUM", "DEFAULT 'tuan'", "Loại: tuan, thang, cuoi_ky, tong_ket"),
            (5, "mo_ta", "TEXT", "NULL", "Mô tả đợt nộp"),
            (6, "start_at", "DATETIME", "NOT NULL", "Thời điểm bắt đầu nhận nộp"),
            (7, "end_at", "DATETIME", "NOT NULL", "Thời điểm kết thúc nhận nộp"),
            (8, "created_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP", "Thời điểm tạo bản ghi"),
        ],
    },
    {
        "index": "2.13",
        "name": "bai_nop_cua_sinh_vien",
        "title": "Bảng bai_nop_cua_sinh_vien (Bài nộp của sinh viên)",
        "columns": [
            (1, "id", "INT", "PK, NOT NULL, AUTO_INCREMENT", "Mã bài nộp"),
            (2, "slot_id", "INT", "NOT NULL, FK → dot_nop_bao_cao_theo_tuan(id)", "Mã đợt nộp"),
            (3, "ma_sinh_vien", "VARCHAR(20)", "NOT NULL", "Mã sinh viên nộp bài"),
            (4, "file_path", "VARCHAR(512)", "NOT NULL", "Đường dẫn file bài nộp"),
            (5, "original_name", "VARCHAR(255)", "NULL", "Tên file gốc"),
            (6, "mime_type", "VARCHAR(100)", "NULL", "Loại file (MIME type)"),
            (7, "file_size", "INT", "NULL", "Kích thước file (bytes)"),
            (8, "submitted_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP", "Thời điểm nộp bài"),
            (9, "teacher_comment", "TEXT", "NULL", "Nhận xét của giảng viên"),
            (10, "trang_thai", "ENUM", "DEFAULT 'da_nop'", "Trạng thái: da_nop, da_duyet, tu_choi"),
        ],
    },
    {
        "index": "2.14",
        "name": "diem_theo_dot_nop",
        "title": "Bảng diem_theo_dot_nop (Điểm theo đợt nộp)",
        "columns": [
            (1, "id", "INT", "PK, NOT NULL, AUTO_INCREMENT", "Mã bản ghi điểm"),
            (2, "slot_id", "INT", "UNIQUE, NOT NULL, FK → dot_nop_bao_cao_theo_tuan(id)", "Mã đợt nộp"),
            (3, "ma_sinh_vien", "VARCHAR(20)", "UNIQUE, NOT NULL", "Mã sinh viên"),
            (4, "diem_giang_vien", "DECIMAL(4,2)", "NULL", "Điểm giảng viên chấm"),
            (5, "nhan_xet_giang_vien", "TEXT", "NULL", "Nhận xét của giảng viên"),
            (6, "updated_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP ON UPDATE", "Thời điểm cập nhật gần nhất"),
        ],
    },
    {
        "index": "2.15",
        "name": "sinh_vien_thuc_tap",
        "title": "Bảng sinh_vien_thuc_tap (Danh sách SV tham gia đợt)",
        "columns": [
            (1, "id", "INT", "PK, NOT NULL, AUTO_INCREMENT", "Mã bản ghi"),
            (2, "ma_sinh_vien", "VARCHAR(20)", "UNIQUE, NOT NULL", "Mã sinh viên"),
            (3, "dot_thuc_tap_id", "INT", "UNIQUE, NOT NULL, FK → dot_thuc_tap(id)", "Mã đợt thực tập"),
            (4, "ngay_dang_ky", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP", "Ngày đăng ký tham gia"),
            (5, "trang_thai", "ENUM", "DEFAULT 'dang-ky'", "dang-ky, duoc-phan-cong, hoan-thanh, huy"),
        ],
    },
    {
        "index": "2.16",
        "name": "notifications",
        "title": "Bảng notifications (Thông báo hệ thống)",
        "columns": [
            (1, "id", "INT", "PK, NOT NULL, AUTO_INCREMENT", "Mã thông báo"),
            (2, "account_id", "INT", "NOT NULL", "accounts.id của người nhận"),
            (3, "title", "VARCHAR(255)", "NOT NULL", "Tiêu đề thông báo"),
            (4, "message", "TEXT", "NOT NULL", "Nội dung thông báo"),
            (5, "type", "ENUM", "NOT NULL, DEFAULT 'info'", "Loại: info, success, warning, error"),
            (6, "is_read", "TINYINT(1)", "NOT NULL, DEFAULT 0", "Trạng thái đã đọc (0: chưa, 1: đã đọc)"),
            (7, "action_type", "VARCHAR(100)", "NULL", "Loại hành động liên quan"),
            (8, "receiver_id", "INT", "NULL", "accounts.id của người nhận"),
            (9, "student_id", "INT", "NULL", "sinh_vien.id nếu người nhận là sinh viên"),
            (10, "created_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP", "Thời điểm gửi thông báo"),
        ],
    },
    {
        "index": "2.17",
        "name": "internship_timeline_milestones",
        "title": "Bảng internship_timeline_milestones (Mốc thời gian quy trình)",
        "columns": [
            (1, "id", "INT", "PK, NOT NULL, AUTO_INCREMENT", "Mã mốc thời gian"),
            (2, "dot_thuc_tap_id", "INT", "UNIQUE, NOT NULL, FK → dot_thuc_tap(id)", "Mã đợt thực tập"),
            (3, "moc_code", "ENUM", "UNIQUE, NOT NULL", "Mã mốc: M1, M2, M3, M4, M5, M6"),
            (4, "ten_moc", "VARCHAR(255)", "NOT NULL", "Tên mốc thời gian"),
            (5, "start_at", "DATETIME", "NOT NULL", "Thời điểm bắt đầu mốc"),
            (6, "end_at", "DATETIME", "NOT NULL", "Thời điểm kết thúc mốc"),
            (7, "sort_order", "TINYINT", "NOT NULL", "Thứ tự hiển thị"),
            (8, "owner_roles", "VARCHAR(255)", "NULL", "Vai trò chủ thực hiện"),
            (9, "recipient_roles", "VARCHAR(255)", "NULL", "Vai trò người nhận"),
            (10, "reminder_offsets", "VARCHAR(255)", "NULL", "Cấu hình nhắc nhở trước hạn"),
            (11, "is_required", "TINYINT(1)", "NOT NULL, DEFAULT 1", "Bắt buộc hay không"),
            (12, "is_active", "TINYINT(1)", "NOT NULL, DEFAULT 1", "Đang kích hoạt hay không"),
            (13, "created_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP", "Thời điểm tạo bản ghi"),
            (14, "updated_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP ON UPDATE", "Thời điểm cập nhật gần nhất"),
        ],
    },
    {
        "index": "2.18",
        "name": "internship_workflow_history",
        "title": "Bảng internship_workflow_history (Lịch sử thay đổi trạng thái)",
        "columns": [
            (1, "id", "BIGINT", "PK, NOT NULL, AUTO_INCREMENT", "Mã lịch sử"),
            (2, "entity_type", "ENUM", "NOT NULL", "Đối tượng: dang_ky_thuc_tap_sinh_vien, phan_cong_thuc_tap"),
            (3, "entity_id", "INT", "NOT NULL", "Mã bản ghi đối tượng"),
            (4, "from_status", "VARCHAR(50)", "NULL", "Trạng thái trước khi thay đổi"),
            (5, "to_status", "VARCHAR(50)", "NOT NULL", "Trạng thái sau khi thay đổi"),
            (6, "changed_by_account_id", "INT", "FK → accounts(id)", "Mã tài khoản thực hiện thay đổi"),
            (7, "changed_by_role", "VARCHAR(30)", "NULL", "Vai trò người thực hiện thay đổi"),
            (8, "note", "TEXT", "NULL", "Ghi chú thay đổi"),
            (9, "changed_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP", "Thời điểm thay đổi"),
        ],
    },
    {
        "index": "2.19",
        "name": "deadline_reminders",
        "title": "Bảng deadline_reminders (Nhắc hạn nộp báo cáo)",
        "columns": [
            (1, "id", "INT", "PK, NOT NULL, AUTO_INCREMENT", "Mã nhắc hạn"),
            (2, "submission_period_id", "INT", "UNIQUE, NOT NULL", "Mã đợt nộp liên quan"),
            (3, "student_id", "INT", "UNIQUE, NOT NULL", "Mã sinh viên nhận nhắc"),
            (4, "type", "ENUM", "UNIQUE, NOT NULL", "Loại: report, diary"),
            (5, "reminder_type", "VARCHAR(50)", "UNIQUE, NOT NULL, DEFAULT 'before_24h'", "Loại nhắc: before_24h, ..."),
            (6, "sent_at", "DATETIME", "DEFAULT CURRENT_TIMESTAMP", "Thời điểm đã gửi nhắc"),
            (7, "created_at", "DATETIME", "DEFAULT CURRENT_TIMESTAMP", "Thời điểm tạo bản ghi"),
        ],
    },
    {
        "index": "2.20",
        "name": "zalo_message_queue",
        "title": "Bảng zalo_message_queue (Hàng đợi tin nhắn Zalo)",
        "columns": [
            (1, "id", "INT", "PK, NOT NULL, AUTO_INCREMENT", "Mã tin nhắn"),
            (2, "lecturer_id", "INT", "NULL", "Mã giảng viên liên quan"),
            (3, "student_id", "INT", "UNIQUE, NOT NULL", "Mã sinh viên nhận tin"),
            (4, "phone", "VARCHAR(20)", "NULL", "Số điện thoại nhận tin"),
            (5, "title", "VARCHAR(255)", "NOT NULL", "Tiêu đề tin nhắn"),
            (6, "message", "TEXT", "NOT NULL", "Nội dung tin nhắn"),
            (7, "type", "ENUM", "UNIQUE, NOT NULL", "new_report_period, new_diary_period, deadline_24h_reminder, manual"),
            (8, "related_id", "INT", "UNIQUE", "Mã đối tượng liên quan"),
            (9, "status", "ENUM", "DEFAULT 'pending'", "pending, processing, sent, failed, cancelled"),
            (10, "priority", "INT", "DEFAULT 5", "Mức độ ưu tiên gửi"),
            (11, "scheduled_at", "DATETIME", "DEFAULT CURRENT_TIMESTAMP", "Thời điểm lên lịch gửi"),
            (12, "sent_at", "DATETIME", "NULL", "Thời điểm đã gửi thành công"),
            (13, "failed_reason", "TEXT", "NULL", "Lý do gửi thất bại"),
            (14, "retry_count", "INT", "DEFAULT 0", "Số lần gửi lại"),
            (15, "created_at", "DATETIME", "DEFAULT CURRENT_TIMESTAMP", "Thời điểm tạo bản ghi"),
            (16, "updated_at", "DATETIME", "DEFAULT CURRENT_TIMESTAMP ON UPDATE", "Thời điểm cập nhật gần nhất"),
        ],
    },
]

# ── Màu sắc ──────────────────────────────────────────────────────────────────
COLOR_HEADER_BG  = RGBColor(0x1F, 0x49, 0x7D)   # xanh đậm (header bảng)
COLOR_HEADER_FG  = RGBColor(0xFF, 0xFF, 0xFF)    # trắng
COLOR_CAPTION_FG = RGBColor(0x1F, 0x49, 0x7D)   # xanh cho caption
COLOR_ROW_ODD    = RGBColor(0xDE, 0xEB, 0xF7)    # xanh nhạt (hàng lẻ)


def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(shd)


def set_cell_borders(cell):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"),   "single")
        border.set(qn("w:sz"),    "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "BFBFBF")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def cell_text(cell, text, bold=False, center=False, font_size=10,
              color: RGBColor | None = None):
    cell.text = ""
    para = cell.paragraphs[0]
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = "Times New Roman"
    if color:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def build_document(output_path: str):
    doc = Document()

    # ── Cài đặt trang A4 ─────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ── Tiêu đề tài liệu ─────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("MÔ TẢ CÁC BẢNG CƠ SỞ DỮ LIỆU MYSQL")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"
    run.font.color.rgb = COLOR_HEADER_BG

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run(
        "Hệ thống Quản lý Thực tập và Hợp tác Doanh nghiệp - Khoa CNTT - ĐH Đại Nam"
    )
    sub_run.italic = True
    sub_run.font.size = Pt(11)
    sub_run.font.name = "Times New Roman"

    doc.add_paragraph()  # khoảng trắng

    # ── Vẽ từng bảng ─────────────────────────────────────────────────────────
    col_widths = [Cm(1.2), Cm(4.5), Cm(3.0), Cm(4.5), Cm(5.0)]
    headers    = ["STT", "Tên cột", "Kiểu dữ liệu", "Ràng buộc", "Mô tả"]

    for tbl in TABLES:
        # Caption
        cap_para = doc.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_para.add_run(f"Bảng {tbl['index']}: {tbl['title']}")
        cap_run.bold = True
        cap_run.font.size = Pt(11)
        cap_run.font.name = "Times New Roman"
        cap_run.font.color.rgb = COLOR_CAPTION_FG

        # Tạo bảng Word
        word_tbl = doc.add_table(rows=1 + len(tbl["columns"]), cols=5)
        word_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        word_tbl.style = "Table Grid"

        # Hàng header
        hdr_cells = word_tbl.rows[0].cells
        for j, (hdr, w) in enumerate(zip(headers, col_widths)):
            hdr_cells[j].width = w
            set_cell_bg(hdr_cells[j], COLOR_HEADER_BG)
            set_cell_borders(hdr_cells[j])
            cell_text(hdr_cells[j], hdr, bold=True, center=True,
                      font_size=10, color=COLOR_HEADER_FG)

        # Hàng dữ liệu
        for i, (stt, col_name, dtype, constraint, desc) in enumerate(tbl["columns"]):
            row_cells = word_tbl.rows[i + 1].cells
            row_bg = COLOR_ROW_ODD if i % 2 == 0 else None

            values = [str(stt), col_name, dtype, constraint, desc]
            for j, val in enumerate(values):
                row_cells[j].width = col_widths[j]
                if row_bg:
                    set_cell_bg(row_cells[j], row_bg)
                set_cell_borders(row_cells[j])
                cell_text(row_cells[j], val,
                          center=(j == 0),
                          font_size=10)

        doc.add_paragraph()  # khoảng trắng giữa các bảng

    doc.save(output_path)
    print(f"[OK] Da xuat: {output_path}")


if __name__ == "__main__":
    out = os.path.join(
        os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
        "backend", "docs", "Mo_ta_20_Bang_CSDL_MySQL.docx"
    )
    build_document(out)
