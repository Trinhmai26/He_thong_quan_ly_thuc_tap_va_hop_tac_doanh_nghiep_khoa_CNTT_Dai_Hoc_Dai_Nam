import json
import os
from collections import defaultdict

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
META_PATH = os.path.join(ROOT, "backend", "docs", "csdl_schema_metadata.json")
OUT_PATH = os.path.join(ROOT, "backend", "docs", "Mo_ta_CSDL_MySQL_UTF8.docx")


TABLE_DESCRIPTIONS = {
    "accounts": "Lưu thông tin tài khoản đăng nhập, vai trò và trạng thái hoạt động của người dùng.",
    "admin": "Lưu thông tin quản trị viên liên kết với tài khoản hệ thống.",
    "sinh_vien": "Lưu hồ sơ sinh viên, thông tin học tập, CV, nguyện vọng và trạng thái thực tập.",
    "giang_vien": "Lưu hồ sơ giảng viên, thông tin chuyên môn và số lượng sinh viên hướng dẫn.",
    "doanh_nghiep": "Lưu thông tin doanh nghiệp, người liên hệ, vị trí tuyển dụng và chỉ tiêu nhận thực tập.",
    "dot_thuc_tap": "Lưu thông tin các đợt thực tập, thời gian đăng ký, thời gian thực tập và trạng thái đợt.",
    "phan_cong_thuc_tap": "Lưu thông tin phân công sinh viên với giảng viên hướng dẫn, doanh nghiệp và đợt thực tập.",
    "dang_ky_thuc_tap_sinh_vien": "Lưu hồ sơ đăng ký thực tập, nguyện vọng, thông tin công ty, trạng thái duyệt và phỏng vấn.",
    "bai_nop_cua_sinh_vien": "Lưu thông tin file bài nộp của sinh viên theo từng đợt nộp báo cáo.",
    "bao_cao_thuc_tap": "Lưu nội dung báo cáo thực tập, file đính kèm, nhận xét, điểm và trạng thái xử lý.",
    "dot_nop_bao_cao_theo_tuan": "Lưu các đợt/lịch nộp báo cáo theo tuần hoặc theo yêu cầu của giảng viên.",
    "diem_theo_dot_nop": "Lưu điểm và nhận xét của giảng viên theo từng đợt nộp báo cáo.",
    "notifications": "Lưu thông báo trong hệ thống gửi tới tài khoản hoặc sinh viên.",
    "zalo_message_queue": "Lưu hàng đợi tin nhắn Zalo phục vụ gửi thông báo tuần tự và kiểm soát trạng thái gửi.",
    "deadline_reminders": "Lưu thông tin nhắc hạn nộp báo cáo để tránh gửi trùng thông báo.",
    "internship_timeline_milestones": "Lưu các mốc thời gian trong tiến trình thực tập của từng đợt.",
    "internship_workflow_history": "Lưu lịch sử thay đổi trạng thái trong quy trình thực tập.",
    "sinh_vien_thuc_tap": "Lưu danh sách sinh viên tham gia một đợt thực tập.",
    "tin_tuyen_dung": "Lưu tin tuyển dụng hoặc vị trí thực tập do doanh nghiệp đăng.",
    "ung_tuyen": "Lưu thông tin sinh viên ứng tuyển vào tin tuyển dụng hoặc vị trí thực tập.",
}


COLUMN_DESCRIPTIONS = {
    "id": "Mã định danh duy nhất của bản ghi.",
    "account_id": "Mã tài khoản liên kết với hồ sơ người dùng.",
    "sinh_vien_id": "Mã sinh viên liên kết với bản ghi.",
    "giang_vien_id": "Mã giảng viên liên kết với bản ghi.",
    "doanh_nghiep_id": "Mã doanh nghiệp liên kết với bản ghi.",
    "dot_thuc_tap_id": "Mã đợt thực tập liên kết với bản ghi.",
    "phan_cong_id": "Mã phân công thực tập liên kết với báo cáo.",
    "slot_id": "Mã đợt nộp báo cáo liên kết với bài nộp hoặc điểm.",
    "tin_tuyen_dung_id": "Mã tin tuyển dụng liên kết với ứng tuyển.",
    "receiver_id": "Mã tài khoản người nhận thông báo.",
    "student_id": "Mã sinh viên liên quan tới thông báo hoặc hàng đợi gửi tin.",
    "changed_by_account_id": "Mã tài khoản thực hiện thay đổi trạng thái.",
    "created_at": "Ngày tạo bản ghi.",
    "updated_at": "Ngày cập nhật bản ghi.",
    "email": "Địa chỉ email.",
    "role": "Vai trò người dùng trong hệ thống.",
    "trang_thai": "Trạng thái nghiệp vụ của bản ghi.",
    "status": "Trạng thái xử lý hoặc trạng thái hoạt động.",
    "is_active": "Trạng thái kích hoạt của tài khoản/bản ghi.",
    "password_hash": "Mật khẩu đã được mã hóa.",
    "file_path": "Đường dẫn lưu file trong hệ thống.",
    "cv_path": "Đường dẫn file CV của sinh viên.",
    "zalo_user_id": "Mã người dùng Zalo đã liên kết.",
}


def humanize(name: str) -> str:
    return name.replace("_", " ")


def infer_desc(col: str, table: str, comment: str | None) -> str:
    if comment and str(comment).strip():
        return str(comment).strip()

    c = col.lower()
    if c in COLUMN_DESCRIPTIONS:
        return COLUMN_DESCRIPTIONS[c]
    if c.startswith("ma_"):
        return "Mã " + humanize(c[3:]) + " dùng để định danh nghiệp vụ."
    if c in ("ho_ten", "full_name") or "ho_ten" in c:
        return "Họ tên của người dùng hoặc đối tượng liên quan."
    if c.startswith("ten_") or c == "ten":
        return "Tên " + humanize(c.replace("ten_", "")) + "."
    if "email" in c:
        return "Địa chỉ email liên hệ."
    if "phone" in c or "dien_thoai" in c or "sdt" in c:
        return "Số điện thoại liên hệ."
    if "dia_chi" in c:
        return "Địa chỉ liên hệ hoặc địa chỉ đơn vị."
    if "ngay" in c or "date" in c:
        return "Thông tin ngày phục vụ nghiệp vụ."
    if "time" in c or "gio" in c:
        return "Thông tin thời gian phục vụ nghiệp vụ."
    if "deadline" in c or "han" in c:
        return "Thời hạn xử lý hoặc hạn nộp."
    if "diem" in c:
        return "Điểm đánh giá hoặc kết quả chấm."
    if "noi_dung" in c or "content" in c or "message" in c:
        return "Nội dung văn bản hoặc thông báo."
    if "mo_ta" in c or "description" in c:
        return "Mô tả chi tiết của bản ghi."
    if "ghi_chu" in c or "note" in c:
        return "Ghi chú bổ sung."
    if "ly_do" in c or "reason" in c:
        return "Lý do xử lý hoặc thay đổi trạng thái."
    if "nhan_xet" in c or "feedback" in c or "danh_gia" in c:
        return "Nhận xét, phản hồi hoặc đánh giá."
    if c.endswith("_id"):
        return "Mã liên kết tới bảng hoặc thực thể liên quan."
    return "Thuộc tính " + humanize(c) + " của bảng " + table + "."


def set_run_font(run, size=12, bold=None, italic=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def apply_cell_font(cell, size=12, bold=False):
    for para in cell.paragraphs:
        for run in para.runs:
            set_run_font(run, size=size, bold=bold)


def set_cell_text(cell, text, align=None, bold=False):
    cell.text = "" if text is None else str(text)
    if align is not None:
        cell.paragraphs[0].alignment = align
    apply_cell_font(cell, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def constraint_text(col, unique_cols, fks_by_col):
    parts = []
    key = col.get("COLUMN_KEY") or ""
    full = f"{col['TABLE_NAME']}.{col['COLUMN_NAME']}"
    if key == "PRI":
        parts.append("PK")
    if full in unique_cols:
        parts.append("UK")
    if col.get("IS_NULLABLE") == "NO":
        parts.append("NOT NULL")
    if col.get("EXTRA"):
        extra = str(col["EXTRA"]).upper().replace("DEFAULT_GENERATED", "DEFAULT")
        parts.append(extra)
    if col.get("COLUMN_DEFAULT") is not None:
        parts.append("DEFAULT")
    fk = fks_by_col.get(full)
    if fk:
        parts.append(f"FK -> {fk['REFERENCED_TABLE_NAME']}({fk['REFERENCED_COLUMN_NAME']})")
    if key == "MUL" and not fk:
        parts.append("INDEX")
    return ", ".join(parts)


def add_caption(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    set_run_font(run, italic=True)


def main():
    with open(META_PATH, "r", encoding="utf-8") as f:
        meta = json.load(f)

    fks_by_col = {f"{fk['TABLE_NAME']}.{fk['COLUMN_NAME']}": fk for fk in meta["fks"]}
    unique_cols = {f"{u['TABLE_NAME']}.{u['COLUMN_NAME']}" for u in meta["uniques"]}
    cols_by_table = defaultdict(list)
    for col in meta["columns"]:
        cols_by_table[col["TABLE_NAME"]].append(col)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)

    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("2.2.1. Xác định thực thể và quan hệ")
    set_run_font(run, size=14, bold=True)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    text = (
        f"Cơ sở dữ liệu của hệ thống sử dụng MySQL với database {meta['database']}. "
        f"Qua quá trình đọc cấu trúc trực tiếp từ hệ quản trị cơ sở dữ liệu, hệ thống hiện có {len(meta['tables'])} bảng chính. "
        "Các bảng được thiết kế nhằm lưu trữ thông tin tài khoản, sinh viên, giảng viên, doanh nghiệp, "
        "đợt thực tập, đăng ký thực tập, phân công, báo cáo, thông báo và các nghiệp vụ hỗ trợ."
    )
    set_run_font(para.add_run(text))

    add_caption(doc, "Bảng 2.1. Danh sách các thực thể trong cơ sở dữ liệu")
    entity_table = doc.add_table(rows=1, cols=4)
    entity_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    entity_table.style = "Table Grid"
    for idx, header in enumerate(["STT", "Tên bảng", "Mô tả", "Số cột"]):
        set_cell_text(entity_table.rows[0].cells[idx], header, WD_ALIGN_PARAGRAPH.CENTER, True)
    for i, table_meta in enumerate(meta["tables"], 1):
        name = table_meta["TABLE_NAME"]
        row = entity_table.add_row().cells
        set_cell_text(row[0], i, WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row[1], name)
        set_cell_text(row[2], table_meta.get("TABLE_COMMENT") or TABLE_DESCRIPTIONS.get(name, f"Bảng {name} trong hệ thống."))
        set_cell_text(row[3], len(cols_by_table[name]), WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()
    add_caption(doc, "Bảng 2.2. Quan hệ khóa ngoại giữa các bảng")
    rel_table = doc.add_table(rows=1, cols=6)
    rel_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rel_table.style = "Table Grid"
    for idx, header in enumerate(["STT", "Bảng con", "Cột FK", "Bảng cha", "Cột tham chiếu", "Quan hệ"]):
        set_cell_text(rel_table.rows[0].cells[idx], header, WD_ALIGN_PARAGRAPH.CENTER, True)
    for i, fk in enumerate(meta["fks"], 1):
        row = rel_table.add_row().cells
        set_cell_text(row[0], i, WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row[1], fk["TABLE_NAME"])
        set_cell_text(row[2], fk["COLUMN_NAME"])
        set_cell_text(row[3], fk["REFERENCED_TABLE_NAME"])
        set_cell_text(row[4], fk["REFERENCED_COLUMN_NAME"])
        set_cell_text(row[5], "1 - N", WD_ALIGN_PARAGRAPH.CENTER)

    for idx, table_meta in enumerate(meta["tables"], 3):
        name = table_meta["TABLE_NAME"]
        doc.add_paragraph()
        add_caption(doc, f"Bảng 2.{idx}. Bảng {name}")
        table = doc.add_table(rows=1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        headers = ["STT", "Tên cột", "Kiểu dữ liệu", "Ràng buộc", "Mô tả"]
        for h_idx, header in enumerate(headers):
            set_cell_text(table.rows[0].cells[h_idx], header, WD_ALIGN_PARAGRAPH.CENTER, True)

        for stt, col in enumerate(cols_by_table[name], 1):
            row = table.add_row().cells
            set_cell_text(row[0], stt, WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(row[1], col["COLUMN_NAME"])
            set_cell_text(row[2], str(col["COLUMN_TYPE"]).upper())
            set_cell_text(row[3], constraint_text(col, unique_cols, fks_by_col))
            set_cell_text(row[4], infer_desc(col["COLUMN_NAME"], name, col.get("COLUMN_COMMENT")))

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
