#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Standalone report validation script (PDF + DOCX).
Called directly by Node.js via child_process — no Flask service required.

Usage:
    python validate_report.py <file_path> <student_name> <original_filename>

Output (JSON to stdout):
    {
        "expected_name":   "Nguyễn Thị Ngọc Ánh",
        "extracted_name":  "Phan Văn Đằng",
        "filename_match":  true,
        "content_match":   false,
        "is_match":        false,
        "message":         "..."
    }
"""

import sys
import io
import json
import os
import re
import unicodedata
import difflib

# Force UTF-8 stdout so Vietnamese characters survive Windows cp1252 terminal
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

_HERE = os.path.dirname(os.path.abspath(__file__))
CV_ANALYZER_DIR = os.path.normpath(os.path.join(_HERE, '..', 'backend', 'cv_analyzer'))
sys.path.insert(0, CV_ANALYZER_DIR)


# ---------------------------------------------------------------------------
# Helpers (mirrored from cv_processing.py but self-contained)
# ---------------------------------------------------------------------------

def normalize_vi(text: str) -> str:
    nfkd = unicodedata.normalize('NFKD', text.lower())
    return ''.join(c for c in nfkd if not unicodedata.combining(c))


VI_UP = 'A-ZÁÀẢÃẠĂẮẰẲẴẶÂẤẦẨẪẬĐÉÈẺẼẸÊẾỀỂỄỆÍÌỈĨỊÓÒỎÕỌÔỐỒỔỖỘƠỚỜỞỠỢÚÙỦŨỤƯỨỪỬỮỰÝỲỶỸỴ'
VI_LO = 'a-záàảãạăắằẳẵặâấầẩẫậđéèẻẽẹêếềểễệíìỉĩịóòỏõọôốồổỗộơớờởỡợúùủũụưứừửữựýỳỷỹỵ'
VI_CH = VI_UP + VI_LO

SKIP_WORDS = {
    # CV section headers
    'kỹ năng', 'kinh nghiệm', 'học vấn', 'giáo dục', 'liên hệ', 'ngôn ngữ',
    'sở thích', 'mục tiêu', 'giới thiệu', 'chứng chỉ', 'thông tin', 'cá nhân',
    'skills', 'education', 'experience', 'contact', 'objective', 'summary',
    'profile', 'references', 'projects', 'hoat dong', 'thanh tich',
    # Report / institutional headers
    'bao cao', 'thực tập', 'thuc tap', 'khoa cntt', 'dai hoc', 'dai nam',
    'truong dai', 'nhan xet', 'nhận xét',
    'bo giao duc', 'bộ giáo dục', 'cong hoa xa hoi', 'cộng hòa xã hội',
    'doc lap tu do', 'độc lập tự do', 'hanh phuc', 'hạnh phúc',
    'phong giao duc', 'phòng giáo dục', 'so giao duc', 'sở giáo dục',
    'truong trung', 'trường trung', 'viet nam', 'việt nam',
    'giao vien huong dan', 'giảng viên hướng dẫn', 'giang vien',
    'sinh vien thuc hien', 'sinh viên thực hiện',
    'khoa cong nghe', 'khoa công nghệ',
}

_MIXED_RE = re.compile(
    rf'(?<![{VI_CH}])([{VI_UP}][{VI_LO}]+'
    rf'(?:\s+[{VI_UP}][{VI_LO}]+){{1,4}})(?![{VI_CH}])',
    re.UNICODE
)
_ALLCAPS_RE = re.compile(
    rf'(?<![{VI_CH}])([{VI_UP}]+'
    rf'(?:\s+[{VI_UP}]+){{1,4}})(?![{VI_CH}])',
    re.UNICODE
)


def _is_skip(candidate: str) -> bool:
    return normalize_vi(candidate.strip()) in SKIP_WORDS


def find_name_in_lines(lines: list, max_lines: int = 30) -> str | None:
    """Try to extract a Vietnamese person name from the first `max_lines` lines."""
    for line in lines[:max_lines]:
        stripped = line.strip()
        if not stripped or _is_skip(stripped):
            continue
        # ALL-CAPS first (names often set in large caps at the top)
        for m in _ALLCAPS_RE.finditer(stripped):
            words = m.group(1).split()
            if 2 <= len(words) <= 5 and not _is_skip(m.group(1)):
                return m.group(1).title()
        # Mixed-case
        for m in _MIXED_RE.finditer(stripped):
            words = m.group(1).split()
            if 2 <= len(words) <= 5 and not _is_skip(m.group(1)):
                return m.group(1)
    return None


def compare_names(expected: str, found: str | None) -> tuple[bool, float]:
    """Return (is_match, similarity) between expected and found name."""
    if not found:
        return False, 0.0
    ne = normalize_vi(expected.strip())
    nf = normalize_vi(found.strip())
    if ne == nf:
        return True, 1.0
    sim = difflib.SequenceMatcher(None, ne, nf).ratio()
    return sim >= 0.82, round(sim, 3)


# ---------------------------------------------------------------------------
# Extractors
# ---------------------------------------------------------------------------

def extract_text_pdf(file_path: str) -> str:
    """Extract text from PDF with Y-sorted blocks (name at visual top comes first)."""
    import fitz
    doc = fitz.open(file_path)
    result = []
    for i, page in enumerate(doc):
        blocks = page.get_text('blocks')
        text_blocks = sorted(
            [b for b in blocks if b[6] == 0 and b[4].strip()],
            key=lambda b: (round(b[1] / 10) * 10, b[0])
        )
        result.extend(b[4].strip() for b in text_blocks)
        if i >= 1:          # first 2 pages are enough for name search
            break
    return '\n'.join(result)


def extract_text_docx(file_path: str) -> str:
    """Extract text from DOCX; header/body paragraphs sorted naturally."""
    import docx
    doc = docx.Document(file_path)
    lines = []
    # Headers appear first in the document body
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            lines.append(t)
    # Also check tables (some reports put names inside tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    lines.append(t)
    return '\n'.join(lines)


def extract_text(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.pdf':
        return extract_text_pdf(file_path)
    if ext in ('.docx', '.doc'):
        return extract_text_docx(file_path)
    return ''


# ---------------------------------------------------------------------------
# Name in filename
# ---------------------------------------------------------------------------

def check_name_in_filename(filename: str, expected_name: str) -> bool:
    """Check whether the student's name (or normalized form) appears in the filename."""
    norm_filename = normalize_vi(os.path.splitext(filename)[0])
    norm_name = normalize_vi(expected_name.strip())

    # Exact substring
    if norm_name in norm_filename:
        return True

    # Token overlap: all name tokens must appear
    tokens = [t for t in norm_name.split() if len(t) > 1]
    if tokens and all(t in norm_filename for t in tokens):
        return True

    # Fuzzy for short filenames
    sim = difflib.SequenceMatcher(None, norm_name, norm_filename).ratio()
    return sim >= 0.75


# ---------------------------------------------------------------------------
# Main validation
# ---------------------------------------------------------------------------

def validate_report(file_path: str, student_name: str, original_filename: str) -> dict:
    if not os.path.isfile(file_path):
        return _err(f"File không tồn tại: {file_path}", student_name)

    # 1. Filename check
    filename_match = check_name_in_filename(original_filename, student_name)

    # 2. Content extraction
    try:
        text = extract_text(file_path)
    except Exception as e:
        return _err(f"Không đọc được nội dung file: {e}", student_name)

    if not text.strip():
        return {
            'expected_name': student_name,
            'extracted_name': None,
            'filename_match': filename_match,
            'content_match': False,
            'is_match': False,
            'message': 'Không đọc được nội dung file (file rỗng hoặc bị mã hóa)'
        }

    # 3. Find name in content (search first 40 lines)
    lines = [l for l in text.split('\n') if l.strip()]
    extracted_name = find_name_in_lines(lines, max_lines=40)

    # 4. Fallback: exact substring in first 3000 chars
    norm_student = normalize_vi(student_name)
    norm_head = normalize_vi(text[:3000])
    if norm_student in norm_head:
        content_match = True
    elif extracted_name:
        content_match, _ = compare_names(student_name, extracted_name)
    else:
        content_match = False

    is_match = content_match  # filename_match is informational only

    # Build message
    if is_match:
        msg = 'Tên sinh viên khớp với nội dung báo cáo'
    else:
        found_str = f"'{extracted_name}'" if extracted_name else 'không tìm thấy tên'
        msg = (
            f"Tên trong báo cáo ({found_str}) không khớp với "
            f"tên sinh viên ('{student_name}')"
        )

    return {
        'expected_name': student_name,
        'extracted_name': extracted_name,
        'filename_match': filename_match,
        'content_match': content_match,
        'is_match': is_match,
        'message': msg
    }


def _err(msg: str, student_name: str) -> dict:
    return {
        'expected_name': student_name,
        'extracted_name': None,
        'filename_match': False,
        'content_match': False,
        'is_match': False,
        'message': msg,
        'error': msg
    }


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main():
    if len(sys.argv) < 4:
        print(json.dumps(
            _err('Usage: validate_report.py <file_path> <student_name> <original_filename>', ''),
            ensure_ascii=False
        ))
        sys.exit(1)

    file_path = sys.argv[1]
    student_name = sys.argv[2].strip()
    original_filename = sys.argv[3].strip()

    result = validate_report(file_path, student_name, original_filename)
    print(json.dumps(result, ensure_ascii=False))


if __name__ == '__main__':
    main()
