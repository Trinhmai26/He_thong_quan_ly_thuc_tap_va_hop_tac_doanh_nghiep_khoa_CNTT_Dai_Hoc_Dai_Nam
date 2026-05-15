from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

OUTPUT_PATH = (
    r"c:\doan\He_thong_quan_ly_thuc_tap_va_hop_tac_doanh_nghiep_khoa_CNTT_Dai_Hoc_Dai_Nam"
    r"\backend\uploads\word\nhatkythuctap.docx"
)

PAGE_WIDTH_CM = 21.0
LEFT_MARGIN_CM = 3.0
RIGHT_MARGIN_CM = 2.0
CONTENT_WIDTH_CM = PAGE_WIDTH_CM - LEFT_MARGIN_CM - RIGHT_MARGIN_CM


def cm_to_twips(value_cm):
    return int(round(value_cm / 2.54 * 1440))


def set_run_font(run, size=12, bold=False, italic=False, underline=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), "Times New Roman")


def set_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=0, line_spacing=1.3):
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def add_text(paragraph, text, size=12, bold=False, italic=False, underline=False):
    run = paragraph.add_run(text)
    set_run_font(run, size, bold, italic, underline)
    return run


def set_cell_margins(cell, top=35, start=70, bottom=35, end=70):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        margin = tc_mar.find(qn(f"w:{name}"))
        if margin is None:
            margin = OxmlElement(f"w:{name}")
            tc_mar.append(margin)
        margin.set(qn("w:w"), str(value))
        margin.set(qn("w:type"), "dxa")


def get_or_add_tbl_pr(table):
    tbl_pr = table._tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)
    return tbl_pr


def set_table_width(table, column_widths_cm, align=WD_TABLE_ALIGNMENT.CENTER):
    table.alignment = align
    table.autofit = False
    tbl_pr = get_or_add_tbl_pr(table)

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(cm_to_twips(sum(column_widths_cm))))
    tbl_w.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    for index, width_cm in enumerate(column_widths_cm):
        width_twips = cm_to_twips(width_cm)
        table.columns[index].width = Cm(width_cm)
        for cell in table.columns[index].cells:
            cell.width = Cm(width_cm)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width_twips))
            tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, visible=True):
    tbl_pr = get_or_add_tbl_pr(table)
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{side}"))
        if border is None:
            border = OxmlElement(f"w:{side}")
            borders.append(border)
        border.set(qn("w:val"), "single" if visible else "nil")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")


def set_row_height(row, height_cm):
    tr_pr = row._tr.get_or_add_trPr()
    tr_height = tr_pr.find(qn("w:trHeight"))
    if tr_height is None:
        tr_height = OxmlElement("w:trHeight")
        tr_pr.append(tr_height)
    tr_height.set(qn("w:val"), str(cm_to_twips(height_cm)))
    tr_height.set(qn("w:hRule"), "atLeast")


def shade_cell(cell, fill="D9D9D9"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)


def add_dotted_line(document, labels, tab_positions_cm):
    paragraph = document.add_paragraph()
    set_paragraph(paragraph, before=0, after=0, line_spacing=1.3)
    for position in tab_positions_cm:
        paragraph.paragraph_format.tab_stops.add_tab_stop(
            Cm(position), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.DOTS
        )
    for label in labels:
        add_text(paragraph, label)
        add_text(paragraph, "\t")


doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

section = doc.sections[0]
section.page_width = Cm(PAGE_WIDTH_CM)
section.page_height = Cm(29.7)
section.left_margin = Cm(LEFT_MARGIN_CM)
section.right_margin = Cm(RIGHT_MARGIN_CM)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)

header_table = doc.add_table(rows=1, cols=2)
set_table_width(header_table, [7.2, 8.8])
set_table_borders(header_table, visible=False)

left_cell = header_table.cell(0, 0)
right_cell = header_table.cell(0, 1)

p = left_cell.paragraphs[0]
set_paragraph(p, WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.15)
add_text(p, "BỘ GIÁO DỤC VÀ ĐÀO TẠO")
p = left_cell.add_paragraph()
set_paragraph(p, WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.15)
add_text(p, "TRƯỜNG ĐẠI HỌC ĐẠI NAM", bold=True, underline=True)

p = right_cell.paragraphs[0]
set_paragraph(p, WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.15)
add_text(p, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", bold=True)
p = right_cell.add_paragraph()
set_paragraph(p, WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.15)
add_text(p, "Độc lập - Tự do - Hạnh phúc", bold=True, underline=True)

p = doc.add_paragraph()
set_paragraph(p, WD_ALIGN_PARAGRAPH.RIGHT, before=2, after=8, line_spacing=1.15)
add_text(p, "Hà Nội, ngày …. tháng … năm 202..", italic=True)

p = doc.add_paragraph()
set_paragraph(p, WD_ALIGN_PARAGRAPH.CENTER, before=0, after=8, line_spacing=1.2)
add_text(p, "NHẬT KÝ THỰC TẬP", size=14, bold=True)

add_dotted_line(doc, ["Họ và tên sinh viên:"], [CONTENT_WIDTH_CM])
add_dotted_line(doc, ["Ngày sinh:", "Nơi sinh:"], [5.1, CONTENT_WIDTH_CM])
add_dotted_line(doc, ["MSV:", "Khóa:", "Lớp:"], [4.4, 9.0, CONTENT_WIDTH_CM])
add_dotted_line(doc, ["Ngành đào tạo:"], [CONTENT_WIDTH_CM])
add_dotted_line(doc, ["Cơ sở thực tập tốt nghiệp:"], [CONTENT_WIDTH_CM])

p = doc.add_paragraph()
set_paragraph(p, after=4, line_spacing=1.0)

week_rows = [f"Tuần {week} (từ ....đến...)" for week in range(1, 9)]
table = doc.add_table(rows=1 + len(week_rows), cols=4)
set_table_width(table, [3.7, 6.2, 4.1, 2.0])
set_table_borders(table, visible=True)

headers = ["Thời gian", "Nội dung thực tập", "Kết quả thực hiện", "Ghi chú"]
for index, header in enumerate(headers):
    cell = table.cell(0, index)
    shade_cell(cell)
    p = cell.paragraphs[0]
    set_paragraph(p, WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.1)
    add_text(p, header, bold=True)
set_row_height(table.rows[0], 0.7)

for row_index, label in enumerate(week_rows, start=1):
    row = table.rows[row_index]
    set_row_height(row, 0.85)
    for cell in row.cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell)
    p = table.cell(row_index, 0).paragraphs[0]
    set_paragraph(p, WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.1)
    add_text(p, label)

p = doc.add_paragraph()
set_paragraph(p, WD_ALIGN_PARAGRAPH.LEFT, before=6, after=0, line_spacing=1.3)
p.paragraph_format.tab_stops.add_tab_stop(Cm(CONTENT_WIDTH_CM), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.DOTS)
add_text(p, "Nhận xét của giảng viên hướng dẫn:")
add_text(p, "\t")

for _ in range(3):
    p = doc.add_paragraph()
    set_paragraph(p, WD_ALIGN_PARAGRAPH.LEFT, before=0, after=0, line_spacing=1.3)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(CONTENT_WIDTH_CM), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.DOTS)
    add_text(p, "\t")

p = doc.add_paragraph()
set_paragraph(p, after=8, line_spacing=1.0)

signature_table = doc.add_table(rows=2, cols=2)
set_table_width(signature_table, [8.0, 8.0])
set_table_borders(signature_table, visible=False)

signature_titles = ["Giảng viên hướng dẫn", "Họ và tên sinh viên"]
signature_notes = ["(Ký và ghi rõ họ tên)", "(Ký và ghi rõ họ tên)"]
for index in range(2):
    p = signature_table.cell(0, index).paragraphs[0]
    set_paragraph(p, WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.15)
    add_text(p, signature_titles[index], bold=True)
    p = signature_table.cell(1, index).paragraphs[0]
    set_paragraph(p, WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.15)
    add_text(p, signature_notes[index], italic=True)

doc.save(OUTPUT_PATH)
print(f"Đã cập nhật file: {OUTPUT_PATH}")
