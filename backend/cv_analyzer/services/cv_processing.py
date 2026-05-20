import fitz  # PyMuPDF
import re
import cv2
from ultralytics import YOLO
import pytesseract
from PIL import Image
import os
import unicodedata
import difflib
import logging
from concurrent.futures import ThreadPoolExecutor
import numpy as np
import tempfile

logger = logging.getLogger(__name__)

# Định nghĩa các từ khóa cho từng mục
SECTION_KEYWORDS = {
    'Kỹ năng': ['Kỹ năng', 'Skills', 'Skill'],
    'Kinh nghiệm': ['Kinh nghiệm', 'Experience', 'Work Experience'],
    'Học vấn': ['Học vấn', 'Education', 'Academic Background']
}

# Mapping label YOLO sang tên section
LABEL_TO_SECTION = {
    3: 'Contact',
    4: 'Education',
    5: 'Skills'
}

YOLO_MODEL = None
YOLO_MODEL_PATH = None


def get_yolo_model(model_path):
    global YOLO_MODEL, YOLO_MODEL_PATH
    normalized_model_path = os.path.abspath(model_path)
    if YOLO_MODEL is None or YOLO_MODEL_PATH != normalized_model_path:
        YOLO_MODEL = YOLO(normalized_model_path)
        YOLO_MODEL_PATH = normalized_model_path
    return YOLO_MODEL

def split_sentences(text):
    """Tách văn bản thành các câu."""
    sentences = re.split(r'(?<=[.!?])\s+', text)
    return [s.strip() for s in sentences if len(s.strip()) > 0]

def extract_sections(text, section_keywords):
    """Tách các mục chính trong CV dựa trên từ khóa."""
    results = {}
    text = text.replace('\r', '\n')
    all_keywords = []
    keyword_to_section = {}

    for section, keywords in section_keywords.items():
        for kw in keywords:
            all_keywords.append(re.escape(kw))
            keyword_to_section[kw.lower()] = section

    pattern = '|'.join(all_keywords)
    matches = list(re.finditer(pattern, text, flags=re.IGNORECASE))

    for i, match in enumerate(matches):
        matched_keyword = match.group().lower()
        section_title = keyword_to_section.get(matched_keyword, matched_keyword)
        start = match.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        content = text[start:end].strip()
        results[section_title] = content

    return results

def ocr_crop(crop):
    pil_crop = Image.fromarray(cv2.cvtColor(crop, cv2.COLOR_BGR2RGB))
    return pytesseract.image_to_string(pil_crop, lang='eng+fra+vie').strip()

def detect_layout_and_ocr(image_path, label_to_section, model_path):
    image = cv2.imread(image_path)
    model = get_yolo_model(model_path)
    results = model(image)
    sections = {}
    crops = []
    section_names = []
    for r in results[0].boxes:
        x1, y1, x2, y2 = map(int, r.xyxy[0])
        if (x2-x1)*(y2-y1) < 5000:
            continue
        label = int(r.cls[0])
        section_name = label_to_section.get(label, f"section_{label}")
        crop = image[y1:y2, x1:x2]
        crops.append(crop)
        section_names.append(section_name)
    # Song song OCR các vùng
    with ThreadPoolExecutor() as executor:
        ocr_results = list(executor.map(ocr_crop, crops))
    for name, text in zip(section_names, ocr_results):
        if name not in sections or len(text) > len(sections[name]):
            sections[name] = text
    return sections

def process_pdf(file_path, model_path):
    import fitz
    doc = fitz.open(file_path)
    full_text = ""
    for page in doc:
        pix = page.get_pixmap(dpi=120)  # thử giảm xuống 120
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        # Lưu ảnh ra RAM, chuyển sang numpy array cho YOLO
        img_np = cv2.cvtColor(np.array(img), cv2.COLOR_RGB2BGR)
        # Lưu tạm ra file nếu YOLO không nhận numpy array, nếu nhận thì truyền trực tiếp
        fd, temp_path = tempfile.mkstemp(prefix="cv_page_", suffix=".png")
        os.close(fd)
        cv2.imwrite(temp_path, img_np)
        sections = detect_layout_and_ocr(temp_path, LABEL_TO_SECTION, model_path)
        for section in sections.values():
            full_text += " " + section
        if os.path.exists(temp_path):
            os.remove(temp_path)
    cv_sections = extract_sections(full_text, SECTION_KEYWORDS)
    cv_sentences = split_sentences(full_text)
    return cv_sections, cv_sentences

def pdf_to_images_with_fitz(pdf_path, output_folder=".", prefix="page"):
    doc = fitz.open(pdf_path)
    image_paths = []
    for i, page in enumerate(doc):
        pix = page.get_pixmap(dpi=300)
        img_path = f"{output_folder}/{prefix}_{i+1}.png"
        pix.save(img_path)
        image_paths.append(img_path)
    return image_paths


# ==============================================================================
# Raw text extraction (no YOLO / OCR — fast, for name validation)
# ==============================================================================

def extract_raw_text(file_path: str, max_pages: int = 2) -> str:
    """
    Trích xuất văn bản thô từ PDF bằng PyMuPDF (không cần YOLO/OCR).
    Các block được sort theo vị trí Y để tên ứng viên ở đầu trang lên trước.

    Args:
        file_path:  Đường dẫn tới file PDF.
        max_pages:  Số trang tối đa cần đọc (mặc định 2 — đủ để tìm tên).

    Returns:
        Chuỗi văn bản đã ghép từ các block text.
    """
    ext = os.path.splitext(file_path)[1].lower()
    if ext in ('.docx', '.doc'):
        return _extract_raw_text_docx(file_path)

    try:
        doc = fitz.open(file_path)
    except Exception as e:
        logger.error("extract_raw_text: cannot open '%s': %s", file_path, e)
        return ''

    result_lines: list[str] = []
    for i, page in enumerate(doc):
        if i >= max_pages:
            break
        try:
            blocks = page.get_text('blocks')
            # Sort by row (Y bucketed to 10px) then X — so visual top-left comes first
            text_blocks = sorted(
                [b for b in blocks if b[6] == 0 and b[4].strip()],
                key=lambda b: (round(b[1] / 10) * 10, b[0])
            )
            result_lines.extend(b[4].strip() for b in text_blocks)
        except Exception as e:
            logger.warning("extract_raw_text: page %d error: %s", i, e)
            continue

    return '\n'.join(result_lines)


def _extract_raw_text_docx(file_path: str) -> str:
    """Extract plain text from DOCX file."""
    try:
        import docx
        doc = docx.Document(file_path)
        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        lines.append(cell.text.strip())
        return '\n'.join(lines)
    except Exception as e:
        logger.error("_extract_raw_text_docx: %s", e)
        return ''


# ==============================================================================
# Name extraction helpers
# ==============================================================================

# Vietnamese + ASCII uppercase/lowercase character classes
_VI_UP = 'A-ZÁÀẢÃẠĂẮẰẲẴẶÂẤẦẨẪẬĐÉÈẺẼẸÊẾỀỂỄỆÍÌỈĨỊÓÒỎÕỌÔỐỒỔỖỘƠỚỜỞỠỢÚÙỦŨỤƯỨỪỬỮỰÝỲỶỸỴ'
_VI_LO = 'a-záàảãạăắằẳẵặâấầẩẫậđéèẻẽẹêếềểễệíìỉĩịóòỏõọôốồổỗộơớờởỡợúùủũụưứừửữựýỳỷỹỵ'
_VI_CH = _VI_UP + _VI_LO

# Pattern: 2–5 Title-Case or ALL-CAPS Vietnamese words
_MIXED_NAME_RE = re.compile(
    rf'(?<![{_VI_CH}])'
    rf'([{_VI_UP}][{_VI_LO}]+'
    rf'(?:\s+[{_VI_UP}][{_VI_LO}]+){{1,4}})'
    rf'(?![{_VI_CH}])',
    re.UNICODE,
)
_ALLCAPS_NAME_RE = re.compile(
    rf'(?<![{_VI_CH}])'
    rf'([{_VI_UP}]{{2,}}'
    rf'(?:\s+[{_VI_UP}]{{2,}}){{1,4}})'
    rf'(?![{_VI_CH}])',
    re.UNICODE,
)

# Words that are section headers / institutional names — not person names
_SKIP_WORDS: frozenset[str] = frozenset({
    'ky nang', 'kinh nghiem', 'hoc van', 'giao duc', 'lien he', 'ngon ngu',
    'so thich', 'muc tieu', 'gioi thieu', 'chung chi', 'thong tin', 'ca nhan',
    'skills', 'education', 'experience', 'contact', 'objective', 'summary',
    'profile', 'references', 'projects',
    'bao cao', 'thuc tap', 'khoa cntt', 'dai hoc', 'dai nam',
    'truong dai', 'nhan xet',
    'bo giao duc', 'cong hoa xa hoi', 'doc lap tu do', 'hanh phuc',
    'viet nam', 'giao vien huong dan', 'giang vien',
    'sinh vien thuc hien', 'khoa cong nghe',
})


def _norm_vi(text: str) -> str:
    """Normalize Vietnamese text: remove diacritics, lowercase, collapse spaces."""
    text = text.replace('đ', 'd').replace('Đ', 'D')
    nfkd = unicodedata.normalize('NFKD', text)
    no_accent = ''.join(c for c in nfkd if not unicodedata.combining(c))
    return re.sub(r'\s+', ' ', no_accent.lower()).strip()


def _is_skip(candidate: str) -> bool:
    return _norm_vi(candidate.strip()) in _SKIP_WORDS


def _looks_like_name(line: str) -> bool:
    """
    Kiểm tra một dòng có thể là tên người không:
    - 2–5 từ
    - Sau khi normalize, mỗi từ chỉ gồm chữ cái [a-z] (1–15 ký tự)
    - Ít nhất một từ dài >= 2 ký tự (loại trừ dòng gồm toàn chữ đơn)
    - Không trong skip list
    """
    if _is_skip(line):
        return False
    norm = _norm_vi(line)
    words = norm.split()
    return (
        2 <= len(words) <= 5
        and all(re.match(r'^[a-z]{1,15}$', w) for w in words)
        and any(len(w) >= 2 for w in words)
    )


def extract_name_from_text(text: str, max_lines: int = 40) -> str | None:
    """
    Tìm tên ứng viên từ nội dung CV (dùng normalize-first approach).

    Chiến lược:
      1. Duyệt max_lines dòng đầu.
      2. Bỏ qua dòng chứa số / ký tự đặc biệt.
      3. Normalize dòng → kiểm tra có phải 2-5 từ thuần chữ cái không.
         (Xử lý: NGUYỄN THANH BÌNH, Nguyen Van A, Nguyễn Thị B, …)
      4. Bỏ qua skip words (tên trường, tiêu đề mục CV, …).

    Returns:
        Tên gốc từ CV (chưa normalize), hoặc None nếu không tìm thấy.
    """
    lines = [ln.strip() for ln in text.split('\n') if ln.strip()]

    for line in lines[:max_lines]:
        # Bỏ dòng có số, ký tự đặc biệt (số điện thoại, email, ngày tháng…)
        if re.search(r'[0-9@/:.()\[\]{},;!?=+*&^%$#~|<>_]', line):
            continue

        if _looks_like_name(line):
            logger.debug("extract_name_from_text: found '%s'", line)
            return line.strip()

    logger.debug("extract_name_from_text: no name found in first %d lines", max_lines)
    return None


# ==============================================================================
# Name comparison
# ==============================================================================

def normalize_name(text: str) -> str:
    """
    Normalize Vietnamese text for name matching.
    đ/Đ must be replaced BEFORE NFKD because they don't decompose under Unicode.
    """
    if not text or not text.strip():
        return ''
    # đ/Đ doesn't decompose under NFKD — handle explicitly first
    text = text.replace('đ', 'd').replace('Đ', 'D')
    text = unicodedata.normalize('NFKD', text)
    text = ''.join(c for c in text if not unicodedata.combining(c))
    text = text.lower()
    # Replace non-alpha chars (digits, punctuation) with space so tokens stay separated
    text = re.sub(r'[^a-z\s]', ' ', text)
    return re.sub(r'\s+', ' ', text).strip()


def _token_set_ratio(a: str, b: str) -> float:
    """Jaccard + sorted-token SequenceMatcher similarity."""
    sa, sb = set(a.split()), set(b.split())
    if not sa or not sb:
        return 0.0
    jaccard = len(sa & sb) / len(sa | sb)
    seq = difflib.SequenceMatcher(
        None, ' '.join(sorted(sa)), ' '.join(sorted(sb))
    ).ratio()
    return max(jaccard, seq)


def compare_names(expected: str, extracted: str, threshold: float = 0.80) -> dict:
    """
    So sánh hai tên sau khi normalize.

    Returns dict with keys: match, similarity, method,
    normalized_expected, normalized_extracted.
    """
    if not expected or not extracted:
        return {'match': False, 'similarity': 0.0, 'method': 'no_match',
                'normalized_expected': normalize_name(expected or ''),
                'normalized_extracted': normalize_name(extracted or '')}

    ne = normalize_name(expected)
    nx = normalize_name(extracted)

    if ne == nx:
        return {'match': True, 'similarity': 1.0, 'method': 'exact',
                'normalized_expected': ne, 'normalized_extracted': nx}

    if ne in nx or nx in ne:
        return {'match': True, 'similarity': 0.95, 'method': 'substring',
                'normalized_expected': ne, 'normalized_extracted': nx}

    seq_r = difflib.SequenceMatcher(None, ne, nx).ratio()
    tok_r = _token_set_ratio(ne, nx)
    best  = round(max(seq_r, tok_r), 3)
    method = 'fuzzy' if seq_r >= tok_r else 'token_set'
    match  = best >= threshold

    logger.info("compare_names | %s → %s | ratio=%.3f (seq=%.3f tok=%.3f)",
                ne, nx, best, seq_r, tok_r)
    return {'match': match, 'similarity': best,
            'method': method if match else 'no_match',
            'normalized_expected': ne, 'normalized_extracted': nx}


def check_filename_match(filename: str, expected_name: str) -> bool:
    """Kiểm tra tên sinh viên có xuất hiện trong tên file không."""
    if not filename or not expected_name:
        return False
    stem = os.path.splitext(filename)[0]
    norm_fn = normalize_name(stem)
    norm_nm = normalize_name(expected_name)
    if norm_nm in norm_fn:
        return True
    tokens = [t for t in norm_nm.split() if len(t) > 1]
    if tokens and all(t in norm_fn for t in tokens):
        return True
    return difflib.SequenceMatcher(None, norm_nm, norm_fn).ratio() >= 0.72


# ==============================================================================
# High-level validate functions (called by validate_cv_upload.py)
# ==============================================================================

def validate_upload(file_path: str, student_name: str,
                    original_filename: str = '',
                    student_code: str = '',
                    student_email: str = '') -> dict:
    """
    Full CV upload validation — content-first, multi-strategy matching.
    Verifies by: name (4 strategies) → student_code → student_email.
    Security: only content match grants access; filename is informational only.
    """
    logger.info("=== validate_upload START ===")
    logger.info("file='%s'  student='%s'  code='%s'  email='%s'",
                os.path.basename(file_path), student_name, student_code, student_email)

    # ── File exists ───────────────────────────────────────────────────────────
    if not os.path.isfile(file_path):
        return _vld_err(f"File không tồn tại: {file_path}", student_name)

    # ── Extract text (read more pages for completeness) ───────────────────────
    try:
        text = extract_raw_text(file_path, max_pages=5)
    except Exception as e:
        logger.error("extract_raw_text failed: %s", e)
        return _vld_err(f"Không đọc được nội dung CV: {e}", student_name)

    text_len = len(text.strip()) if text else 0
    logger.info("extractedTextLength: %d", text_len)
    if text_len > 0:
        logger.info("extractedText (first 2000):\n%s", text[:2000])
    else:
        logger.warning("PDF không trích xuất được text — empty or image-only PDF")

    if not text or text_len < 10:
        logger.warning("REJECT: cannot read CV content")
        fn_match = check_filename_match(original_filename, student_name)
        return {
            'expected_name': student_name, 'extracted_name': None,
            'filename_match': fn_match, 'content_match': False,
            'is_match': False, 'similarity': 0.0,
            'message': (
                'Không đọc được nội dung CV (file rỗng hoặc chỉ chứa ảnh scan). '
                'Vui lòng upload CV dạng PDF có thể copy text.'
            ),
        }

    # ── filename match (informational) ────────────────────────────────────────
    filename_match = check_filename_match(original_filename, student_name)

    # ── Normalize student name ─────────────────────────────────────────────────
    norm_student      = normalize_name(student_name)
    norm_student_ns   = norm_student.replace(' ', '')        # no-space variant
    student_tokens    = [t for t in norm_student.split() if len(t) >= 2]

    logger.info("studentName          : '%s'", student_name)
    logger.info("normalizedStudentName: '%s'", norm_student)
    logger.info("studentTokens        : %s",  student_tokens)

    # ── Normalize full CV text ─────────────────────────────────────────────────
    norm_text    = normalize_name(text)
    norm_text_ns = norm_text.replace(' ', '')

    logger.info("normalizedText (first 500): %s", norm_text[:500])

    # ══════════════════════════════════════════════════════════════════════════
    # STRATEGY 1: Direct normalized substring — fastest, most reliable
    # ══════════════════════════════════════════════════════════════════════════
    s1 = bool(norm_student and norm_student in norm_text)
    logger.info("S1 directSubstring    : %s  ('%s' in text)", s1, norm_student)
    if s1:
        logger.info("ACCEPTED via S1")
        return _content_match_result(student_name, student_name, filename_match, 1.0)

    # ══════════════════════════════════════════════════════════════════════════
    # STRATEGY 2: No-space variant — handles PDF that merges words or splits
    #             across lines (e.g. "trinhthiyenmai" in "trinhthiyenmai26")
    # ══════════════════════════════════════════════════════════════════════════
    s2 = bool(norm_student_ns and len(norm_student_ns) >= 6 and norm_student_ns in norm_text_ns)
    logger.info("S2 noSpaceMatch       : %s  ('%s' in textNoSpace)", s2, norm_student_ns)
    if s2:
        logger.info("ACCEPTED via S2")
        return _content_match_result(student_name, student_name, filename_match, 0.95)

    # ══════════════════════════════════════════════════════════════════════════
    # STRATEGY 3: All-token match — every word of name appears in text
    #             (handles newline-separated names in extracted text)
    # ══════════════════════════════════════════════════════════════════════════
    cv_word_set   = set(norm_text.split())
    # Use EXACT word match (not substring) to prevent "thi" matching in "thiet", "mai" in "email"
    matched_toks  = [t for t in student_tokens if t in cv_word_set]
    s3 = len(student_tokens) >= 2 and len(matched_toks) == len(student_tokens)
    logger.info("S3 allTokensWholeWord : %s  matched=%s / total=%s",
                s3, matched_toks, student_tokens)
    if s3:
        logger.info("ACCEPTED via S3")
        return _content_match_result(student_name, student_name, filename_match, 0.90)

    # ══════════════════════════════════════════════════════════════════════════
    # STRATEGY 4: Extract candidate name → fuzzy compare
    # ══════════════════════════════════════════════════════════════════════════
    extracted_name = extract_name_from_text(text, max_lines=40)
    logger.info("S4 extractedCandidateName: '%s'", extracted_name)

    if extracted_name:
        cmp = compare_names(student_name, extracted_name, threshold=0.80)
        s4_match  = cmp['match']
        similarity = cmp['similarity']
        logger.info("S4 fuzzyMatch=%s  similarity=%.3f", s4_match, similarity)
    else:
        s4_match  = False
        similarity = 0.0
        logger.info("S4: Không tìm thấy tên sinh viên trong extractedText")

    if s4_match:
        logger.info("ACCEPTED via S4")
        return _content_match_result(student_name, extracted_name, filename_match, similarity)

    # ══════════════════════════════════════════════════════════════════════════
    # STRATEGY 5: Student code in raw text (strongest machine-readable signal)
    # ══════════════════════════════════════════════════════════════════════════
    s5 = bool(student_code and student_code.strip() and student_code.strip() in text)
    logger.info("S5 studentCodeInText  : %s  (code='%s')", s5, student_code)
    if s5:
        logger.info("ACCEPTED via S5 (student code found in CV)")
        return _content_match_result(student_name, None, filename_match, 1.0)

    # ══════════════════════════════════════════════════════════════════════════
    # STRATEGY 6: Student email in raw text
    # ══════════════════════════════════════════════════════════════════════════
    s6 = bool(
        student_email and student_email.strip()
        and student_email.strip().lower() in text.lower()
    )
    logger.info("S6 studentEmailInText : %s  (email='%s')", s6, student_email)
    if s6:
        logger.info("ACCEPTED via S6 (student email found in CV)")
        return _content_match_result(student_name, None, filename_match, 1.0)

    # ── All strategies failed ─────────────────────────────────────────────────
    if extracted_name:
        msg = (
            f"Tên ứng viên trong CV ('{extracted_name}') "
            f"không khớp với tài khoản upload ('{student_name}')"
        )
    elif text_len > 0:
        msg = (
            'Không tìm thấy họ tên, mã sinh viên hoặc email trong nội dung CV. '
            'Vui lòng upload CV chứa thông tin cá nhân của bạn.'
        )
    else:
        msg = (
            'Không đọc được nội dung CV. '
            'Vui lòng upload CV dạng PDF có thể copy text (không phải ảnh scan).'
        )

    logger.info("REJECTED | S1=%s S2=%s S3=%s S4=%s S5=%s S6=%s",
                s1, s2, s3, s4_match, s5, s6)
    logger.info("=== validate_upload END ===")

    return {
        'expected_name':  student_name,
        'extracted_name': extracted_name,
        'filename_match': filename_match,
        'content_match':  False,
        'is_match':       False,
        'similarity':     round(similarity, 3),
        'message':        msg,
    }


def _content_match_result(student_name: str, extracted_name: str,
                           filename_match: bool, similarity: float) -> dict:
    return {
        'expected_name':  student_name,
        'extracted_name': extracted_name,
        'filename_match': filename_match,
        'content_match':  True,
        'is_match':       True,
        'similarity':     round(similarity, 3),
        'message':        'Tên ứng viên trong CV khớp với tài khoản upload',
    }


def _vld_err(msg: str, student_name: str) -> dict:
    return {
        'expected_name':  student_name,
        'extracted_name': None,
        'filename_match': False,
        'content_match':  False,
        'is_match':       False,
        'similarity':     0.0,
        'message':        msg,
        'error':          msg,
    }


def validate_student_name_in_cv(cv_path: str, student_name: str) -> dict:
    """
    Simple validation: extract text → find name → compare.
    Used by validate_cv_name.py (legacy entry point).

    Returns:
        {"isMatch": bool, "nameInCV": str|None, "similarity": float, "message": str}
    """
    result = validate_upload(cv_path, student_name, original_filename='')
    return {
        'isMatch':   result['is_match'],
        'nameInCV':  result['extracted_name'],
        'similarity': result['similarity'],
        'message':   result['message'],
    }